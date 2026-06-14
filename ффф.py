import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import psycopg2
from datetime import datetime, timedelta
import hashlib
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

DB_NAME = "aaa"
DB_USER = "postgres"
DB_PASSWORD = "Neren2020"
DB_HOST = "localhost"
DB_PORT = "5432"


def get_db_connection():
    return psycopg2.connect(
        dbname=DB_NAME,
        user=DB_USER,
        password=DB_PASSWORD,
        host=DB_HOST,
        port=DB_PORT
    )


def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()


class MainApp:
    def __init__(self, user_id, username, role, full_name):
        self.user_id = user_id
        self.username = username
        self.role = role
        self.full_name = full_name
        self.pending_transfers = []

        self.root = tk.Tk()
        self.root.title(f"Школьный склад - {full_name} ({role})")
        self.root.geometry("1400x800")

        self.create_menu()
        self.show_dashboard()
        self.check_pending_transfers()

        self.root.mainloop()

    def create_menu(self):
        self.left_frame = tk.Frame(self.root, width=240, bg='#2c3e50')
        self.left_frame.pack(side=tk.LEFT, fill=tk.Y)

        tk.Label(self.left_frame, text="МЕНЮ", font=('Arial', 16, 'bold'), bg='#2c3e50', fg='white').pack(pady=20)

        menu_items = [
            ("📊 Главная", self.show_dashboard),
            ("📦 Расходные материалы", self.show_materials),
            ("🪑 Основные средства", self.show_inventory),
            ("📋 Движения", self.show_movements),
            ("👥 Сотрудники", self.show_persons),
            ("📄 Акт оприходования", self.show_receipt_acts),
            ("📑 Акт списания", self.show_write_off_acts),
            ("⚠️ Акт о расхождениях", self.show_discrepancy_acts),
            ("📈 Отчёты", self.show_reports),
        ]

        if self.role == 'admin':
            menu_items.append(("🔐 Пользователи", self.show_users))
            menu_items.append(("👥 Управление комиссией", self.manage_commission))
            menu_items.append(("📜 Логи", self.show_logs))

        # Словарь для хранения кнопок по ключам
        self.menu_buttons = {}

        for text, command in menu_items:
            # Создаём уникальный ключ для каждой кнопки
            key = text.replace(" ", "_").replace("📊", "").replace("📦", "").replace("🪑", "").replace("📋", "").replace(
                "👥", "").replace("📄", "").replace("📑", "").replace("⚠️", "").replace("📈", "").replace("🔐", "").replace(
                "👥", "").replace("📜", "")
            key = key.strip()

            btn = tk.Button(self.left_frame, text=text,
                            command=lambda t=text, c=command: self.on_menu_click(c, t),
                            width=22, height=2, bg='#34495e', fg='white', bd=0, anchor='w', padx=10)
            btn.pack(pady=3)
            self.menu_buttons[key] = btn

        # Кнопка выхода (не выделяется при нажатии)
        tk.Button(self.left_frame, text="🚪 Выход", command=self.logout, width=22, height=2,
                  bg='#e74c3c', fg='white', bd=0).pack(pady=20)

        self.content_frame = tk.Frame(self.root, bg='#ecf0f1')
        self.content_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

    def reset_menu_colors(self):
        """Сбрасывает цвет всех кнопок меню на исходный"""
        for key, btn in self.menu_buttons.items():
            btn.config(bg='#34495e', fg='white')

    def highlight_menu_button(self, button_text):
        """Выделяет конкретную кнопку меню зелёным цветом"""
        self.reset_menu_colors()
        # Ищем кнопку по тексту
        for btn in self.menu_buttons.values():
            if btn.cget('text') == button_text:
                btn.config(bg='#1abc9c', fg='white')
                break

    def on_menu_click(self, command, button_text):
        """Обработчик клика по меню - выделяет кнопку и выполняет команду"""
        self.highlight_menu_button(button_text)
        command()

    def clear_content(self):
        for widget in self.content_frame.winfo_children():
            widget.destroy()

    def check_pending_transfers(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, name, location, responsible_person_id, transfer_pending, new_location, new_person_id
            FROM inventory_items 
            WHERE transfer_pending = TRUE
        """)
        self.pending_transfers = cur.fetchall()
        cur.close()
        conn.close()

    def show_dashboard(self):
        self.highlight_menu_button("📊 Главная")
        self.clear_content()

        tk.Label(self.content_frame, text="Панель управления", font=('Arial', 20, 'bold'),
                 bg='#ecf0f1', fg='#2c3e50').pack(pady=20)

        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute("SELECT COUNT(*) FROM materials")
        materials_count = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM inventory_items WHERE status = 'in_use'")
        inventory_count = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM materials WHERE min_quantity > 0 AND actual_quantity <= min_quantity")
        low_stock = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM movements WHERE operation_date >= CURRENT_DATE - INTERVAL '30 days'")
        movements_month = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM notifications_view")
        discrepancies_count = cur.fetchone()[0]

        cur.close()
        conn.close()

        stats_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        stats_frame.pack(pady=20)

        stats = [
            ("📦 Расходных материалов", materials_count),
            ("🪑 Основных средств", inventory_count),
            ("⚠️ Мало на складе", low_stock, "#e74c3c" if low_stock > 0 else "#27ae60"),
            ("📋 Операций за месяц", movements_month)
        ]

        for i, stat in enumerate(stats):
            color = stat[2] if len(stat) > 2 else "#3498db"
            frame = tk.Frame(stats_frame, bg='white', relief=tk.RAISED, bd=2)
            frame.grid(row=0, column=i, padx=15, pady=10, sticky='nsew')
            tk.Label(frame, text=stat[0], font=('Arial', 12), bg='white').pack(pady=10)
            tk.Label(frame, text=str(stat[1]), font=('Arial', 28, 'bold'), bg='white', fg=color).pack(pady=10)

        if discrepancies_count > 0 and self.role in ['admin', 'storekeeper']:
            warning_frame = tk.Frame(self.content_frame, bg='#e74c3c', relief=tk.RAISED, bd=2)
            warning_frame.pack(pady=20, padx=20, fill=tk.X)
            tk.Label(warning_frame, text=f"⚠️ ВНИМАНИЕ! Обнаружено {discrepancies_count} расхождений!",
                     font=('Arial', 14, 'bold'), bg='#e74c3c', fg='white').pack(pady=10)
            tk.Button(warning_frame, text="Сформировать акт о расхождениях",
                      command=self.create_discrepancy_act, bg='white', fg='#e74c3c', font=('Arial', 12)).pack(pady=5)

        if low_stock > 0:
            tk.Label(self.content_frame, text=f"⚠️ {low_stock} материалов имеют остаток ниже минимального!",
                     font=('Arial', 12), bg='#ecf0f1', fg='#e74c3c').pack(pady=10)

        if self.pending_transfers and self.role in ['admin', 'storekeeper']:
            transfer_frame = tk.Frame(self.content_frame, bg='#f39c12', relief=tk.RAISED, bd=2)
            transfer_frame.pack(pady=20, padx=20, fill=tk.X)
            tk.Label(transfer_frame, text=f"📦 ВНИМАНИЕ! {len(self.pending_transfers)} основных средств ожидают оформления акта перехода!",
                     font=('Arial', 14, 'bold'), bg='#f39c12', fg='white').pack(pady=10)
            tk.Button(transfer_frame, text="Сформировать акт перехода",
                      command=self.show_transfer_act_dialog, bg='white', fg='#f39c12', font=('Arial', 12)).pack(pady=5)

    def manage_commission(self):
        if self.role != 'admin':
            messagebox.showerror("Доступ запрещён", "Только администратор может управлять комиссией")
            return
        self.highlight_menu_button("👥 Управление комиссией")
        dialog = tk.Toplevel(self.root)
        dialog.title("Управление комиссией по списанию")
        dialog.geometry("500x500")
        dialog.configure(bg='white')
        dialog.resizable(False, False)

        tk.Label(dialog, text="Управление комиссией по списанию", font=('Arial', 14, 'bold'), bg='white').pack(pady=15)

        current_commission = self.get_commission()
        current_frame = tk.LabelFrame(dialog, text="Текущая комиссия", font=('Arial', 11, 'bold'), bg='white')
        current_frame.pack(fill=tk.X, padx=20, pady=10)

        if current_commission:
            tk.Label(current_frame, text=f"Председатель: {current_commission[1]}", font=('Arial', 10), bg='white', anchor='w').pack(anchor='w', padx=10, pady=2)
            tk.Label(current_frame, text=f"Члены: {current_commission[2]}", font=('Arial', 10), bg='white', anchor='w', wraplength=400).pack(anchor='w', padx=10, pady=2)
        else:
            tk.Label(current_frame, text="Комиссия не назначена", font=('Arial', 10, 'italic'), bg='white', fg='gray').pack(pady=10)

        tk.Button(dialog, text="📝 Сформировать новую комиссию", command=lambda: self.show_new_commission_dialog(dialog),
                  bg='#27ae60', fg='white', font=('Arial', 11), width=25).pack(pady=20)

    def show_new_commission_dialog(self, parent_dialog):
        parent_dialog.destroy()

        dialog = tk.Toplevel(self.root)
        dialog.title("Формирование новой комиссии")
        dialog.geometry("500x450")
        dialog.configure(bg='white')
        dialog.resizable(False, False)

        tk.Label(dialog, text="Формирование новой комиссии", font=('Arial', 14, 'bold'), bg='white').pack(pady=15)

        persons = self.get_persons()

        tk.Label(dialog, text="Председатель комиссии:", font=('Arial', 11, 'bold'), bg='white').pack(anchor='w', padx=20, pady=5)
        chairman_combo = ttk.Combobox(dialog, values=persons, width=50)
        chairman_combo.pack(padx=20, pady=5)

        tk.Label(dialog, text="Члены комиссии (3 человека):", font=('Arial', 11, 'bold'), bg='white').pack(anchor='w', padx=20, pady=5)

        member_combos = []
        for i in range(3):
            frame = tk.Frame(dialog, bg='white')
            frame.pack(fill=tk.X, padx=20, pady=3)
            tk.Label(frame, text=f"Член {i+1}:", font=('Arial', 10), bg='white', width=12).pack(side=tk.LEFT)
            combo = ttk.Combobox(frame, values=persons, width=45)
            combo.pack(side=tk.LEFT, padx=5)
            member_combos.append(combo)

        def save_commission():
            chairman = chairman_combo.get().strip()
            members = [m.get().strip() for m in member_combos if m.get().strip()]

            if not chairman:
                messagebox.showerror("Ошибка", "Выберите председателя комиссии")
                return

            if len(members) < 3:
                messagebox.showerror("Ошибка", "Выберите 3 членов комиссии")
                return

            members_str = ', '.join(members[:3])

            if messagebox.askyesno("Подтверждение", "При формировании новой комиссии старая будет расформирована. Продолжить?"):
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("DELETE FROM commissions WHERE commission_type = 'spisaniya'")
                cur.execute("""
                    INSERT INTO commissions (order_number, order_date, commission_type, chairman, members, purpose)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (f"№{datetime.now().year}-01", datetime.now().date(), 'spisaniya', chairman, members_str, 'Комиссия по списанию'))
                conn.commit()
                cur.close()
                conn.close()
                self.log_action(f"Обновлена комиссия: председатель {chairman}, члены {members_str}")
                messagebox.showinfo("Успех", "Комиссия сохранена")
                dialog.destroy()

        tk.Button(dialog, text="Сохранить комиссию", command=save_commission, bg='#27ae60', fg='white', width=20).pack(pady=20)

    def get_locations(self, item_type='both'):
        conn = get_db_connection()
        cur = conn.cursor()
        if item_type == 'both':
            cur.execute("SELECT name FROM locations ORDER BY name")
        else:
            cur.execute("SELECT name FROM locations WHERE type = %s OR type = 'both' ORDER BY name", (item_type,))
        locations = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return locations

    def get_suppliers(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT name FROM suppliers ORDER BY name")
        suppliers = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return suppliers

    def get_categories(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT name FROM asset_categories ORDER BY name")
        categories = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return categories

    def get_units(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT short_name FROM units ORDER BY short_name")
        units = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return units

    def get_persons(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT full_name FROM responsible_persons ORDER BY full_name")
        persons = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return persons

    def get_purposes(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT name FROM purposes ORDER BY name")
        purposes = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return purposes

    def get_commission(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT id, chairman, members FROM commissions WHERE commission_type = 'spisaniya' ORDER BY id DESC LIMIT 1")
        commission = cur.fetchone()
        cur.close()
        conn.close()
        return commission

    def get_category_name_by_id(self, category_id):
        if not category_id:
            return None
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT name FROM asset_categories WHERE id = %s", (category_id,))
        res = cur.fetchone()
        cur.close()
        conn.close()
        return res[0] if res else None

    def get_person_name_by_id(self, person_id):
        if not person_id:
            return None
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT full_name FROM responsible_persons WHERE id = %s", (person_id,))
        res = cur.fetchone()
        cur.close()
        conn.close()
        return res[0] if res else None

    def get_unit_name_by_id(self, unit_id):
        if not unit_id:
            return ''
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT short_name FROM units WHERE id = %s", (unit_id,))
        res = cur.fetchone()
        cur.close()
        conn.close()
        return res[0] if res else ''

    def update_existing_items_list(self, item_type, combo):
        """Обновляет список существующих позиций в зависимости от типа"""
        conn = get_db_connection()
        cur = conn.cursor()
        if item_type == 'Материал':
            cur.execute("""
                SELECT id, name, unit_id, price, actual_quantity, accounting_quantity 
                FROM materials 
                ORDER BY name
            """)
        else:
            cur.execute("""
                SELECT id, name, unit_id, price, actual_quantity, accounting_quantity 
                FROM inventory_items 
                WHERE status = 'in_use' 
                ORDER BY name
            """)
        items = cur.fetchall()
        cur.close()
        conn.close()

        item_list = []
        self.existing_items_data = {}
        for item in items:
            display_name = f"{item[1]} (остаток: {item[4]})"
            item_list.append(display_name)
            self.existing_items_data[display_name] = {
                'id': item[0],
                'name': item[1],
                'unit_id': item[2],
                'price': item[3],
                'actual_quantity': item[4],
                'accounting_quantity': item[5]
            }

        item_list.insert(0, "➕ Новое")
        combo['values'] = item_list
        combo.set("➕ Новое")

    def update_write_off_items_list(self, item_type, combo):
        """Обновляет список существующих позиций для акта списания в зависимости от типа"""
        conn = get_db_connection()
        cur = conn.cursor()
        if item_type == 'материалы':
            cur.execute("""
                SELECT id, name, unit_id, price, actual_quantity, accounting_quantity 
                FROM materials 
                WHERE actual_quantity > 0
                ORDER BY name
            """)
        else:
            cur.execute("""
                SELECT id, name, unit_id, price, actual_quantity, accounting_quantity 
                FROM inventory_items 
                WHERE status = 'in_use' AND actual_quantity > 0
                ORDER BY name
            """)
        items = cur.fetchall()
        cur.close()
        conn.close()

        item_list = []
        self.write_off_items_data = {}
        for item in items:
            display_name = f"{item[1]} (остаток: {item[4]})"
            item_list.append(display_name)
            self.write_off_items_data[display_name] = {
                'id': item[0],
                'name': item[1],
                'unit_id': item[2],
                'price': item[3],
                'actual_quantity': item[4],
                'accounting_quantity': item[5]
            }

        # Только существующие позиции, без "➕ Новое"
        combo['values'] = item_list
        if item_list:
            combo.set(item_list[0])

    def print_text_widget(self, text_widget):
        try:
            text_widget.config(state='normal')
            content = text_widget.get('1.0', tk.END)
            print_window = tk.Toplevel()
            print_window.title("Печать")
            text = tk.Text(print_window, font=('Courier', 10))
            text.insert('1.0', content)
            text.pack()
            text_widget.config(state='disabled')
            messagebox.showinfo("Печать", "Нажмите Ctrl+P для печати", parent=print_window)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка печати: {str(e)}")

    def log_action(self, details):
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO action_logs (username, action, object_type, details, created_at)
                VALUES (%s, %s, %s, %s, %s)
            """, (self.username, 'ADMIN_ACTION', 'user', details, datetime.now()))
            conn.commit()
            cur.close()
            conn.close()
        except:
            pass

    def check_and_show_notification(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM notifications_view")
        count = cur.fetchone()[0]
        cur.close()
        conn.close()

        if count > 0 and self.role in ['admin', 'storekeeper']:
            if messagebox.askyesno("Обнаружены расхождения",
                                   f"Обнаружено {count} расхождений. Сформировать акт о расхождениях?"):
                self.create_discrepancy_act()

    def show_discrepancy_notification(self):
        """Показывает всплывающее уведомление о расхождениях"""
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT item_type, item_name, surplus_quantity, shortage_quantity
            FROM notifications_view
            WHERE surplus_quantity != 0 OR shortage_quantity != 0
        """)
        discrepancies = cur.fetchall()
        cur.close()
        conn.close()

        if discrepancies:
            # Формируем сообщение
            msg = "⚠️ ОБНАРУЖЕНЫ РАСХОЖДЕНИЯ!\n\n"
            for disc in discrepancies:
                item_type, name, surplus, shortage = disc
                if surplus > 0:
                    msg += f"• {name} ({item_type}): излишки +{surplus} шт.\n"
                if shortage > 0:
                    msg += f"• {name} ({item_type}): недостача -{shortage} шт.\n"
            msg += "\nСоздать акт о расхождениях?"

            # Показываем всплывающее окно поверх всех окон
            if messagebox.askyesno("ВНИМАНИЕ! Расхождения", msg):
                self.create_discrepancy_act()

    def sort_treeview(self, tree, col, reverse):
        data_list = [(tree.set(child, col), child) for child in tree.get_children('')]
        try:
            data_list.sort(key=lambda x: float(x[0]) if x[0].replace('.', '').replace('-', '').isdigit() else x[0], reverse=reverse)
        except:
            data_list.sort(key=lambda x: x[0].lower(), reverse=reverse)

        for index, (val, child) in enumerate(data_list):
            tree.move(child, '', index)

        tree.heading(col, command=lambda: self.sort_treeview(tree, col, not reverse))

    def show_materials(self):
        self.highlight_menu_button("📦 Расходные материалы")
        self.clear_content()

        tk.Label(self.content_frame, text="Расходные материалы", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        filter_frame = tk.LabelFrame(self.content_frame, text="Фильтры", font=('Arial', 10, 'bold'),
                                     bg='#ecf0f1', fg='#2c3e50')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Место хранения:", bg='#ecf0f1').grid(row=0, column=0, padx=5, pady=5)
        self.material_location_filter = ttk.Combobox(filter_frame, values=[''] + self.get_locations('material'), width=25)
        self.material_location_filter.grid(row=0, column=1, padx=5, pady=5)
        self.material_location_filter.bind('<<ComboboxSelected>>', lambda e: self.load_materials_data())

        btn_frame = tk.Frame(filter_frame, bg='#ecf0f1')
        btn_frame.grid(row=0, column=2, columnspan=2, padx=20, pady=5)
        if self.role in ['admin', 'storekeeper']:
            tk.Button(btn_frame, text="📄 Создать акт оприходования", command=self.create_receipt_act,
                      bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📑 Создать акт списания", command=self.create_write_off_act,
                      bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📥 Выдать", command=self.issue_material, bg='#3498db', fg='white').pack(side=tk.LEFT, padx=5)

        notebook = ttk.Notebook(self.content_frame)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        actual_frame = tk.Frame(notebook, bg='#ecf0f1')
        accounting_frame = tk.Frame(notebook, bg='#ecf0f1')
        notebook.add(actual_frame, text="📝 Фактическое наличие")
        notebook.add(accounting_frame, text="💰 Данные бухгалтерии")

        self.create_materials_actual_tab(actual_frame)
        self.create_materials_accounting_tab(accounting_frame)

    def create_materials_actual_tab(self, parent):
        search_frame = tk.Frame(parent, bg='#ecf0f1')
        search_frame.pack(pady=5, fill=tk.X)

        tk.Label(search_frame, text="🔍 Поиск:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        self.material_search = tk.Entry(search_frame, width=40)
        self.material_search.pack(side=tk.LEFT, padx=5)
        self.material_search.bind('<KeyRelease>', lambda e: self.filter_materials())
        tk.Button(search_frame, text="Очистить", command=self.clear_material_search, bg='#95a5a6', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'name', 'category', 'unit', 'actual_quantity', 'actual_sum', 'min_quantity', 'location')
        self.material_tree = ttk.Treeview(parent, columns=columns, show='headings', height=18)

        headings = {'id': 'ID', 'name': 'Наименование', 'category': 'Категория', 'unit': 'Ед.изм',
                    'actual_quantity': 'Факт.кол-во', 'actual_sum': 'Факт.сумма',
                    'min_quantity': 'Мин.остаток', 'location': 'Место'}

        for col, heading in headings.items():
            self.material_tree.heading(col, text=heading)
            self.material_tree.column(col, width=100 if col != 'name' else 250)
            self.material_tree.heading(col, command=lambda c=col: self.sort_treeview(self.material_tree, c, False))

        self.material_tree.pack(fill=tk.BOTH, expand=True)
        self.load_materials_data()

    def create_materials_accounting_tab(self, parent):
        btn_frame = tk.Frame(parent, bg='#ecf0f1')
        btn_frame.pack(pady=5)

        if self.role in ['admin', 'accountant']:
            tk.Button(btn_frame, text="✏️ Редактировать учётные данные", command=self.edit_material_accounting,
                      bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=5)

        search_frame = tk.Frame(parent, bg='#ecf0f1')
        search_frame.pack(pady=5, fill=tk.X)

        tk.Label(search_frame, text="🔍 Поиск:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        self.material_accounting_search = tk.Entry(search_frame, width=40)
        self.material_accounting_search.pack(side=tk.LEFT, padx=5)
        self.material_accounting_search.bind('<KeyRelease>', lambda e: self.filter_materials_accounting())
        tk.Button(search_frame, text="Очистить", command=self.clear_material_accounting_search, bg='#95a5a6',
                  fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'name', 'category', 'unit', 'accounting_quantity', 'accounting_sum', 'location')
        self.material_accounting_tree = ttk.Treeview(parent, columns=columns, show='headings', height=18)

        headings = {'id': 'ID', 'name': 'Наименование', 'category': 'Категория', 'unit': 'Ед.изм',
                    'accounting_quantity': 'Учёт.кол-во', 'accounting_sum': 'Учёт.сумма', 'location': 'Место'}

        for col, heading in headings.items():
            self.material_accounting_tree.heading(col, text=heading)
            self.material_accounting_tree.column(col, width=100 if col != 'name' else 250)
            self.material_accounting_tree.heading(col, command=lambda c=col: self.sort_treeview(self.material_accounting_tree, c, False))

        self.material_accounting_tree.pack(fill=tk.BOTH, expand=True)
        self.load_materials_accounting_data()

    def load_materials_data(self):
        for item in self.material_tree.get_children():
            self.material_tree.delete(item)

        location_filter = self.material_location_filter.get()

        conn = get_db_connection()
        cur = conn.cursor()

        query = """
            SELECT m.id, m.name, COALESCE(c.name, '') as category, COALESCE(u.short_name, '') as unit,
                   m.actual_quantity, m.actual_sum, m.min_quantity, COALESCE(m.location, '') as location
            FROM materials m
            LEFT JOIN asset_categories c ON m.category_id = c.id
            LEFT JOIN units u ON m.unit_id = u.id
            WHERE 1=1
        """
        params = []

        if location_filter:
            query += " AND m.location = %s"
            params.append(location_filter)

        query += " ORDER BY m.name"

        cur.execute(query, params)
        self.all_materials_data = cur.fetchall()
        cur.close()
        conn.close()

        self.filter_materials()

    def filter_materials(self):
        search_text = self.material_search.get().strip().lower()
        for item in self.material_tree.get_children():
            self.material_tree.delete(item)
        for row in self.all_materials_data:
            if search_text == "" or search_text in str(row[1]).lower():
                self.material_tree.insert('', tk.END, values=row)

    def clear_material_search(self):
        self.material_search.delete(0, tk.END)
        self.filter_materials()

    def load_materials_accounting_data(self):
        for item in self.material_accounting_tree.get_children():
            self.material_accounting_tree.delete(item)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT m.id, m.name, COALESCE(c.name, '') as category, COALESCE(u.short_name, '') as unit,
                   m.accounting_quantity, m.accounting_sum, COALESCE(m.location, '') as location
            FROM materials m
            LEFT JOIN asset_categories c ON m.category_id = c.id
            LEFT JOIN units u ON m.unit_id = u.id
            ORDER BY m.name
        """)
        self.all_materials_accounting_data = cur.fetchall()
        cur.close()
        conn.close()
        self.filter_materials_accounting()

    def filter_materials_accounting(self):
        search_text = self.material_accounting_search.get().strip().lower()
        for item in self.material_accounting_tree.get_children():
            self.material_accounting_tree.delete(item)
        for row in self.all_materials_accounting_data:
            if search_text == "" or search_text in str(row[1]).lower():
                self.material_accounting_tree.insert('', tk.END, values=row)

    def clear_material_accounting_search(self):
        self.material_accounting_search.delete(0, tk.END)
        self.filter_materials_accounting()

    def edit_material_accounting(self):
        selected = self.material_accounting_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите материал")
            return

        values = self.material_accounting_tree.item(selected[0])['values']
        material_id = values[0]
        material_name = values[1]

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT accounting_quantity FROM materials WHERE id = %s", (material_id,))
        current = cur.fetchone()
        cur.close()
        conn.close()

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Редактировать учётные данные: {material_name}")
        dialog.geometry("400x200")

        tk.Label(dialog, text="Количество по данным бухгалтерии:").grid(row=0, column=0, padx=10, pady=10, sticky='w')
        quantity_entry = tk.Entry(dialog, width=30)
        quantity_entry.insert(0, str(current[0]))
        quantity_entry.grid(row=0, column=1, padx=10, pady=10)

        def save():
            try:
                quantity = int(quantity_entry.get())
                if quantity < 0:
                    messagebox.showerror("Ошибка", "Количество не может быть отрицательным")
                    return

                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("UPDATE materials SET accounting_quantity = %s WHERE id = %s", (quantity, material_id))
                conn.commit()
                cur.close()
                conn.close()
                dialog.destroy()
                self.load_materials_data()
                self.load_materials_accounting_data()
                self.check_and_show_notification()
                messagebox.showinfo("Успех", "Данные обновлены")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        tk.Button(dialog, text="Сохранить", command=save, bg='#f39c12', fg='white').grid(row=1, column=0, columnspan=2, pady=20)

    def issue_material(self):
        selected = self.material_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите материал для выдачи")
            return

        material_id = self.material_tree.item(selected[0])['values'][0]
        material_name = self.material_tree.item(selected[0])['values'][1]

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Выдача материала: {material_name}")
        dialog.geometry("400x400")

        persons = self.get_persons()
        purposes = self.get_purposes()

        tk.Label(dialog, text="Количество:").grid(row=0, column=0, padx=10, pady=5, sticky='w')
        quantity_entry = tk.Entry(dialog, width=30)
        quantity_entry.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Кому выдать:").grid(row=1, column=0, padx=10, pady=5, sticky='w')
        person_combo = ttk.Combobox(dialog, values=persons, width=27)
        person_combo.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Цель:").grid(row=2, column=0, padx=10, pady=5, sticky='w')
        purpose_combo = ttk.Combobox(dialog, values=purposes, width=27)
        purpose_combo.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Основание:").grid(row=3, column=0, padx=10, pady=5, sticky='w')
        reason_entry = tk.Entry(dialog, width=30)
        reason_entry.grid(row=3, column=1, padx=10, pady=5)

        def save():
            try:
                quantity = int(quantity_entry.get())
                if quantity <= 0:
                    messagebox.showerror("Ошибка", "Количество должно быть положительным")
                    return

                person_name = person_combo.get()
                if not person_name:
                    messagebox.showerror("Ошибка", "Выберите сотрудника")
                    return

                conn = get_db_connection()
                cur = conn.cursor()

                cur.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (person_name,))
                person = cur.fetchone()
                if not person:
                    messagebox.showerror("Ошибка", "Сотрудник не найден")
                    return
                person_id = person[0]

                purpose_id = None
                if purpose_combo.get():
                    cur.execute("SELECT id FROM purposes WHERE name = %s", (purpose_combo.get(),))
                    purpose = cur.fetchone()
                    if purpose:
                        purpose_id = purpose[0]

                reason = reason_entry.get().strip()

                cur.execute("""
                    INSERT INTO movements (material_id, operation_type, quantity, person_id, purpose_id, reason, created_by)
                    VALUES (%s, 'Выдача', %s, %s, %s, %s, %s)
                """, (material_id, quantity, person_id, purpose_id, reason, self.username))

                cur.execute(
                    "UPDATE materials SET actual_quantity = actual_quantity - %s, accounting_quantity = accounting_quantity - %s WHERE id = %s",
                    (quantity, quantity, material_id))

                conn.commit()
                cur.close()
                conn.close()
                dialog.destroy()
                self.load_materials_data()
                self.load_materials_accounting_data()
                self.check_and_show_notification()
                messagebox.showinfo("Успех", f"Выдано {quantity} ед. материала {material_name}")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        tk.Button(dialog, text="Выдать", command=save, bg='#3498db', fg='white').grid(row=4, column=0, columnspan=2, pady=20)

    def show_inventory(self):
        self.highlight_menu_button("🪑 Основные средства")
        self.clear_content()

        tk.Label(self.content_frame, text="Основные средства", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        filter_frame = tk.LabelFrame(self.content_frame, text="Фильтры", font=('Arial', 10, 'bold'),
                                     bg='#ecf0f1', fg='#2c3e50')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="МОЛ:", bg='#ecf0f1').grid(row=0, column=0, padx=5, pady=5)
        self.inventory_person_filter = ttk.Combobox(filter_frame, values=[''] + self.get_persons(), width=25)
        self.inventory_person_filter.grid(row=0, column=1, padx=5, pady=5)
        self.inventory_person_filter.bind('<<ComboboxSelected>>', lambda e: self.load_inventory_data())

        tk.Label(filter_frame, text="Место нахождения:", bg='#ecf0f1').grid(row=0, column=2, padx=5, pady=5)
        self.inventory_location_filter = ttk.Combobox(filter_frame, values=[''] + self.get_locations('inventory'), width=25)
        self.inventory_location_filter.grid(row=0, column=3, padx=5, pady=5)
        self.inventory_location_filter.bind('<<ComboboxSelected>>', lambda e: self.load_inventory_data())

        btn_frame = tk.Frame(filter_frame, bg='#ecf0f1')
        btn_frame.grid(row=0, column=4, columnspan=2, padx=20, pady=5)
        if self.role in ['admin', 'storekeeper']:
            tk.Button(btn_frame, text="📄 Создать акт оприходования", command=self.create_receipt_act,
                      bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📑 Создать акт списания", command=self.create_write_off_act,
                      bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📦 Переместить", command=self.move_inventory_item, bg='#3498db', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📄 Инвентарная карточка", command=self.show_inventory_card, bg='#9b59b6', fg='white').pack(side=tk.LEFT, padx=5)

        notebook = ttk.Notebook(self.content_frame)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        actual_frame = tk.Frame(notebook, bg='#ecf0f1')
        accounting_frame = tk.Frame(notebook, bg='#ecf0f1')
        notebook.add(actual_frame, text="📝 Фактическое наличие")
        notebook.add(accounting_frame, text="💰 Данные бухгалтерии")

        self.create_inventory_actual_tab(actual_frame)
        self.create_inventory_accounting_tab(accounting_frame)

    def create_inventory_actual_tab(self, parent):
        search_frame = tk.Frame(parent, bg='#ecf0f1')
        search_frame.pack(pady=5, fill=tk.X)

        tk.Label(search_frame, text="🔍 Поиск:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        self.inventory_search = tk.Entry(search_frame, width=40)
        self.inventory_search.pack(side=tk.LEFT, padx=5)
        self.inventory_search.bind('<KeyRelease>', lambda e: self.filter_inventory())
        tk.Button(search_frame, text="Очистить", command=self.clear_inventory_search, bg='#95a5a6', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'inventory_number', 'name', 'category', 'actual_quantity', 'actual_sum', 'location', 'responsible_person')
        self.inventory_tree = ttk.Treeview(parent, columns=columns, show='headings', height=18)

        headings = {'id': 'ID', 'inventory_number': 'Инв.номер', 'name': 'Наименование', 'category': 'Категория',
                    'actual_quantity': 'Факт.кол-во', 'actual_sum': 'Факт.сумма', 'location': 'Место',
                    'responsible_person': 'МОЛ'}

        for col, heading in headings.items():
            self.inventory_tree.heading(col, text=heading)
            self.inventory_tree.column(col, width=100 if col in ['id', 'actual_quantity'] else 150)
            self.inventory_tree.heading(col, command=lambda c=col: self.sort_treeview(self.inventory_tree, c, False))

        self.inventory_tree.pack(fill=tk.BOTH, expand=True)
        self.load_inventory_data()

    def create_inventory_accounting_tab(self, parent):
        btn_frame = tk.Frame(parent, bg='#ecf0f1')
        btn_frame.pack(pady=5)

        if self.role in ['admin', 'accountant']:
            tk.Button(btn_frame, text="✏️ Редактировать учётные данные", command=self.edit_inventory_accounting,
                      bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=5)

        search_frame = tk.Frame(parent, bg='#ecf0f1')
        search_frame.pack(pady=5, fill=tk.X)

        tk.Label(search_frame, text="🔍 Поиск:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        self.inventory_accounting_search = tk.Entry(search_frame, width=40)
        self.inventory_accounting_search.pack(side=tk.LEFT, padx=5)
        self.inventory_accounting_search.bind('<KeyRelease>', lambda e: self.filter_inventory_accounting())
        tk.Button(search_frame, text="Очистить", command=self.clear_inventory_accounting_search, bg='#95a5a6', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'inventory_number', 'name', 'category', 'accounting_quantity', 'accounting_sum', 'location')
        self.inventory_accounting_tree = ttk.Treeview(parent, columns=columns, show='headings', height=18)

        headings = {'id': 'ID', 'inventory_number': 'Инв.номер', 'name': 'Наименование', 'category': 'Категория',
                    'accounting_quantity': 'Учёт.кол-во', 'accounting_sum': 'Учёт.сумма', 'location': 'Место'}

        for col, heading in headings.items():
            self.inventory_accounting_tree.heading(col, text=heading)
            self.inventory_accounting_tree.column(col, width=100 if col in ['id', 'accounting_quantity'] else 150)
            self.inventory_accounting_tree.heading(col, command=lambda c=col: self.sort_treeview(self.inventory_accounting_tree, c, False))

        self.inventory_accounting_tree.pack(fill=tk.BOTH, expand=True)
        self.load_inventory_accounting_data()

    def load_inventory_data(self):
        for item in self.inventory_tree.get_children():
            self.inventory_tree.delete(item)

        person_filter = self.inventory_person_filter.get()
        location_filter = self.inventory_location_filter.get()

        conn = get_db_connection()
        cur = conn.cursor()

        query = """
            SELECT i.id, i.inventory_number, i.name, COALESCE(c.name, '') as category,
                   i.actual_quantity, i.actual_sum, COALESCE(i.location, '') as location,
                   COALESCE(rp.full_name, '') as responsible_person
            FROM inventory_items i
            LEFT JOIN asset_categories c ON i.category_id = c.id
            LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
            WHERE i.status = 'in_use'
        """
        params = []

        if person_filter:
            query += " AND rp.full_name = %s"
            params.append(person_filter)
        if location_filter:
            query += " AND i.location = %s"
            params.append(location_filter)

        query += " ORDER BY i.name"

        cur.execute(query, params)
        self.all_inventory_data = cur.fetchall()
        cur.close()
        conn.close()
        self.filter_inventory()

    def filter_inventory(self):
        search_text = self.inventory_search.get().strip().lower()
        for item in self.inventory_tree.get_children():
            self.inventory_tree.delete(item)
        for row in self.all_inventory_data:
            if search_text == "" or search_text in str(row[2]).lower() or search_text in str(row[1]).lower():
                self.inventory_tree.insert('', tk.END, values=row)

    def clear_inventory_search(self):
        self.inventory_search.delete(0, tk.END)
        self.filter_inventory()

    def load_inventory_accounting_data(self):
        for item in self.inventory_accounting_tree.get_children():
            self.inventory_accounting_tree.delete(item)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT i.id, i.inventory_number, i.name, COALESCE(c.name, '') as category,
                   i.accounting_quantity, i.accounting_sum, COALESCE(i.location, '') as location
            FROM inventory_items i
            LEFT JOIN asset_categories c ON i.category_id = c.id
            WHERE i.status = 'in_use'
            ORDER BY i.name
        """)
        self.all_inventory_accounting_data = cur.fetchall()
        cur.close()
        conn.close()
        self.filter_inventory_accounting()

    def filter_inventory_accounting(self):
        search_text = self.inventory_accounting_search.get().strip().lower()
        for item in self.inventory_accounting_tree.get_children():
            self.inventory_accounting_tree.delete(item)
        for row in self.all_inventory_accounting_data:
            if search_text == "" or search_text in str(row[2]).lower() or search_text in str(row[1]).lower():
                self.inventory_accounting_tree.insert('', tk.END, values=row)

    def clear_inventory_accounting_search(self):
        self.inventory_accounting_search.delete(0, tk.END)
        self.filter_inventory_accounting()

    def edit_inventory_accounting(self):
        selected = self.inventory_accounting_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите объект")
            return

        values = self.inventory_accounting_tree.item(selected[0])['values']
        item_id = values[0]
        item_name = values[2]

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT accounting_quantity FROM inventory_items WHERE id = %s", (item_id,))
        current = cur.fetchone()
        cur.close()
        conn.close()

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Редактировать учётные данные: {item_name}")
        dialog.geometry("400x200")

        tk.Label(dialog, text="Количество по данным бухгалтерии:").grid(row=0, column=0, padx=10, pady=10, sticky='w')
        quantity_entry = tk.Entry(dialog, width=30)
        quantity_entry.insert(0, str(current[0]))
        quantity_entry.grid(row=0, column=1, padx=10, pady=10)

        def save():
            try:
                quantity = int(quantity_entry.get())
                if quantity < 0:
                    messagebox.showerror("Ошибка", "Количество не может быть отрицательным")
                    return

                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("UPDATE inventory_items SET accounting_quantity = %s WHERE id = %s", (quantity, item_id))
                conn.commit()
                cur.close()
                conn.close()
                dialog.destroy()
                self.load_inventory_data()
                self.load_inventory_accounting_data()
                self.check_and_show_notification()
                messagebox.showinfo("Успех", "Данные обновлены")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        tk.Button(dialog, text="Сохранить", command=save, bg='#f39c12', fg='white').grid(row=1, column=0, columnspan=2, pady=20)

    def move_inventory_item(self):
        selected = self.inventory_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите объект для перемещения")
            return

        values = self.inventory_tree.item(selected[0])['values']
        item_id = values[0]
        item_name = values[2]

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Перемещение: {item_name}")
        dialog.geometry("400x300")

        persons = self.get_persons()
        locations = self.get_locations('inventory')

        tk.Label(dialog, text="Новое место нахождения:").grid(row=0, column=0, padx=10, pady=10, sticky='w')
        location_combo = ttk.Combobox(dialog, values=locations, width=27)
        location_combo.grid(row=0, column=1, padx=10, pady=10)

        tk.Label(dialog, text="Новое ответственное лицо:").grid(row=1, column=0, padx=10, pady=10, sticky='w')
        person_combo = ttk.Combobox(dialog, values=persons, width=27)
        person_combo.grid(row=1, column=1, padx=10, pady=10)

        tk.Label(dialog, text="Основание:").grid(row=2, column=0, padx=10, pady=10, sticky='w')
        reason_entry = tk.Entry(dialog, width=30)
        reason_entry.grid(row=2, column=1, padx=10, pady=10)

        def save():
            try:
                new_location = location_combo.get().strip()
                person_name = person_combo.get()
                reason = reason_entry.get().strip()

                conn = get_db_connection()
                cur = conn.cursor()

                if new_location:
                    cur.execute("UPDATE inventory_items SET location = %s WHERE id = %s", (new_location, item_id))

                if person_name:
                    cur.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (person_name,))
                    person = cur.fetchone()
                    if person:
                        cur.execute("UPDATE inventory_items SET responsible_person_id = %s WHERE id = %s",
                                    (person[0], item_id))

                cur.execute("""
                    INSERT INTO movements (inventory_item_id, operation_type, quantity, reason, created_by)
                    VALUES (%s, 'Перемещение', 1, %s, %s)
                """, (item_id, reason, self.username))

                conn.commit()
                cur.close()
                conn.close()
                dialog.destroy()
                self.load_inventory_data()
                self.check_pending_transfers()
                messagebox.showinfo("Успех", "Объект перемещён")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        tk.Button(dialog, text="Переместить", command=save, bg='#3498db', fg='white').grid(row=3, column=0, columnspan=2, pady=20)

    def show_inventory_card(self, item_id=None):
        if item_id is None:
            selected = self.inventory_tree.selection()
            if not selected:
                messagebox.showwarning("Ошибка", "Выберите основное средство")
                return
            item_id = self.inventory_tree.item(selected[0])['values'][0]

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT i.id, i.inventory_number, i.name, i.price, i.actual_quantity, 
                   i.accounting_quantity, i.location, i.status, i.commissioning_date, 
                   i.useful_life_months, i.description, c.name as category_name,
                   rp.full_name as person_name
            FROM inventory_items i
            LEFT JOIN asset_categories c ON i.category_id = c.id
            LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
            WHERE i.id = %s
        """, (item_id,))
        item = cur.fetchone()
        cur.close()
        conn.close()

        if not item:
            messagebox.showerror("Ошибка", "Объект не найден")
            return

        item_id_val = item[0]
        inv_number = item[1] if item[1] else '________'
        name_val = item[2] if item[2] else '________'
        price_val = item[3] if item[3] else 0
        actual_qty = item[4] if item[4] else 0
        accounting_qty = item[5] if item[5] else 0
        location_val = item[6] if item[6] else '________'
        status_val = item[7] if item[7] else 'in_use'
        commissioning_date = item[8] if item[8] else '________'
        useful_life = item[9] if item[9] else '________'
        description_val = item[10] if item[10] else '________'
        category_name = item[11] if item[11] else '________'
        person_name = item[12] if item[12] else '________'

        actual_sum = actual_qty * price_val if actual_qty and price_val else 0
        accounting_sum = accounting_qty * price_val if accounting_qty and price_val else 0

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Инвентарная карточка - {inv_number}")
        dialog.geometry("650x600")
        dialog.configure(bg='white')
        dialog.resizable(False, False)

        can_edit = self.role in ['admin', 'storekeeper']
        can_edit_accounting = self.role == 'admin'

        canvas = tk.Canvas(dialog, bg='white')
        scrollbar = tk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = scrollable_frame

        tk.Label(main_frame, text="ИНВЕНТАРНАЯ КАРТОЧКА", font=('Arial', 14, 'bold'), bg='white').pack(pady=5)

        period_start = commissioning_date if commissioning_date != '________' else '________'
        period_end = datetime.now().strftime('%d.%m.%Y')
        info_frame = tk.Frame(main_frame, bg='white')
        info_frame.pack()
        tk.Label(info_frame, text=f"Период: с {period_start} по {period_end}", font=('Arial', 9), bg='white').pack()
        tk.Label(info_frame, text="ГБОУ «Простоквашинская школа»", font=('Arial', 9), bg='white').pack()
        tk.Label(info_frame, text=f"Подразделение: {location_val}", font=('Arial', 9), bg='white').pack()

        frame_data = tk.Frame(main_frame, bg='white')
        frame_data.pack(fill=tk.X, pady=5)

        fields_left = [
            ("Инв. номер:", inv_number, False),
            ("Наименование:", name_val[:30] + "..." if len(name_val) > 30 else name_val, False),
            ("Состояние:", "Списано" if status_val == 'written_off' else "В эксплуатации", False),
            ("Дата принятия:", commissioning_date, can_edit),
            ("Срок службы:", f"{useful_life} мес." if useful_life != '________' else '________', can_edit),
            ("Стоимость:", f"{price_val:,.2f} руб.", can_edit),
            ("Место:", location_val, can_edit),
        ]

        fields_right = [
            ("Категория:", category_name, False),
            ("МОЛ:", person_name[:25] + "..." if len(person_name) > 25 else person_name, can_edit),
            ("Кол-во (факт):", actual_qty, can_edit),
            ("Сумма (факт):", f"{actual_sum:,.2f}", False),
            ("Кол-во (учёт):", accounting_qty, can_edit_accounting),
            ("Сумма (учёт):", f"{accounting_sum:,.2f}", False),
            ("Описание:", description_val[:40] + "..." if len(description_val) > 40 else description_val, False),
        ]

        left_col = tk.Frame(frame_data, bg='white')
        left_col.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        right_col = tk.Frame(frame_data, bg='white')
        right_col.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)

        entries = {}
        row = 0
        for label, value, editable in fields_left:
            tk.Label(left_col, text=label, font=('Arial', 9, 'bold'), bg='white', anchor='w').grid(row=row, column=0, padx=3, pady=2, sticky='w')
            if editable:
                entry = tk.Entry(left_col, width=20, font=('Arial', 9))
                clean_value = str(value).replace(" руб.", "").replace(" мес.", "")
                entry.insert(0, clean_value)
                entry.grid(row=row, column=1, padx=3, pady=2, sticky='w')
                if label == "Стоимость:":
                    entries['price'] = entry
                elif label == "Дата принятия:":
                    entries['commissioning_date'] = entry
                elif label == "Срок службы:":
                    entries['useful_life_months'] = entry
                elif label == "Место:":
                    entries['location'] = entry
                elif label == "МОЛ:":
                    entries['person_name'] = entry
            else:
                tk.Label(left_col, text=str(value), font=('Arial', 9), bg='white', anchor='w', wraplength=200).grid(row=row, column=1, padx=3, pady=2, sticky='w')
            row += 1

        row = 0
        for label, value, editable in fields_right:
            tk.Label(right_col, text=label, font=('Arial', 9, 'bold'), bg='white', anchor='w').grid(row=row, column=0, padx=3, pady=2, sticky='w')
            if editable and label in ["Кол-во (факт):", "Кол-во (учёт):"]:
                entry = tk.Entry(right_col, width=15, font=('Arial', 9))
                entry.insert(0, str(value))
                entry.grid(row=row, column=1, padx=3, pady=2, sticky='w')
                if label == "Кол-во (факт):":
                    entries['actual_quantity'] = entry
                elif label == "Кол-во (учёт):":
                    entries['accounting_quantity'] = entry
            elif editable and label == "МОЛ:":
                entry = tk.Entry(right_col, width=20, font=('Arial', 9))
                entry.insert(0, str(value))
                entry.grid(row=row, column=1, padx=3, pady=2, sticky='w')
                entries['person_name'] = entry
            else:
                tk.Label(right_col, text=str(value), font=('Arial', 9), bg='white', anchor='w').grid(row=row, column=1, padx=3, pady=2, sticky='w')
            row += 1

        surplus = actual_qty - accounting_qty if actual_qty and accounting_qty else 0
        if surplus > 0:
            disc_text = f"⚠️ ИЗЛИШКИ: +{surplus} шт. (+{surplus * price_val:,.2f} руб.)"
            disc_color = "#e67e22"
        elif surplus < 0:
            disc_text = f"⚠️ НЕДОСТАЧА: {surplus} шт. ({surplus * price_val:,.2f} руб.)"
            disc_color = "#e74c3c"
        else:
            disc_text = "✓ Расхождений нет"
            disc_color = "#27ae60"
        tk.Label(main_frame, text=disc_text, font=('Arial', 10, 'bold'), bg='white', fg=disc_color).pack(pady=5)

        frame_sign = tk.LabelFrame(main_frame, text="Подписи", font=('Arial', 10, 'bold'), bg='white', fg='#2c3e50')
        frame_sign.pack(fill=tk.X, pady=5, padx=5)

        tk.Label(frame_sign, text="Материально ответственное лицо: __________________ (подпись) ______ (дата)",
                 font=('Arial', 9), bg='white').pack(anchor='w', pady=2)
        tk.Label(frame_sign, text="Председатель комиссии: __________________ (подпись) ______ (дата)",
                 font=('Arial', 9), bg='white').pack(anchor='w', pady=2)
        tk.Label(frame_sign, text="Члены комиссии: __________________ (подпись) ______ (дата)",
                 font=('Arial', 9), bg='white').pack(anchor='w', pady=2)
        tk.Label(frame_sign, text=f"Дата: {datetime.now().strftime('%d.%m.%Y')}",
                 font=('Arial', 9), bg='white').pack(anchor='w', pady=2)

        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=10)

        def save_card():
            if not can_edit:
                messagebox.showwarning("Доступ запрещён", "У вас нет прав на редактирование")
                return
            try:
                conn2 = get_db_connection()
                cur2 = conn2.cursor()
                update_fields = []
                update_values = []

                for field_name, entry in entries.items():
                    val = entry.get().strip()
                    if field_name in ['price', 'actual_quantity', 'accounting_quantity', 'useful_life_months']:
                        try:
                            val = float(val) if '.' in val else int(val) if val else 0
                        except:
                            val = 0
                    elif field_name == 'person_name':
                        if val:
                            cur2.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (val,))
                            res = cur2.fetchone()
                            if res:
                                update_fields.append("responsible_person_id = %s")
                                update_values.append(res[0])
                            else:
                                cur2.execute("INSERT INTO responsible_persons (full_name) VALUES (%s) RETURNING id", (val,))
                                update_fields.append("responsible_person_id = %s")
                                update_values.append(cur2.fetchone()[0])
                        continue
                    update_fields.append(f"{field_name} = %s")
                    update_values.append(val)

                if update_fields:
                    update_values.append(item_id_val)
                    cur2.execute(f"UPDATE inventory_items SET {', '.join(update_fields)} WHERE id = %s", tuple(update_values))

                conn2.commit()
                cur2.close()
                conn2.close()
                self.load_inventory_data()
                self.load_inventory_accounting_data()
                messagebox.showinfo("Успех", "Данные сохранены")
                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        def print_card():
            print_window = tk.Toplevel(dialog)
            print_window.title("Печать инвентарной карточки")
            print_window.geometry("650x700")
            print_window.configure(bg='white')

            text_widget = tk.Text(print_window, bg='white', font=('Courier', 10), wrap=tk.WORD)
            text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            content_lines = []
            content_lines.append("=" * 70)
            content_lines.append(" " * 20 + "ИНВЕНТАРНАЯ КАРТОЧКА")
            content_lines.append("=" * 70)
            content_lines.append(f"Инвентарный номер: {inv_number}")
            content_lines.append(f"Наименование: {name_val}")
            content_lines.append(f"Состояние: {'Списано' if status_val == 'written_off' else 'В эксплуатации'}")
            content_lines.append(f"Дата принятия: {commissioning_date}")
            content_lines.append(f"Срок службы: {useful_life} мес." if useful_life != '________' else "Срок службы: ________")
            content_lines.append(f"Стоимость: {price_val:,.2f} руб.")
            content_lines.append(f"Место: {location_val}")
            content_lines.append(f"Категория: {category_name}")
            content_lines.append(f"МОЛ: {person_name}")
            content_lines.append(f"Описание: {description_val}")
            content_lines.append(f"Количество (факт): {actual_qty}")
            content_lines.append(f"Сумма (факт): {actual_sum:,.2f} руб.")
            content_lines.append(f"Количество (учёт): {accounting_qty}")
            content_lines.append(f"Сумма (учёт): {accounting_sum:,.2f} руб.")

            if surplus > 0:
                content_lines.append(f"ИЗЛИШКИ: +{surplus} шт. (+{surplus * price_val:,.2f} руб.)")
            elif surplus < 0:
                content_lines.append(f"НЕДОСТАЧА: {surplus} шт. ({surplus * price_val:,.2f} руб.)")
            else:
                content_lines.append("Расхождений нет")

            content_lines.append("-" * 70)
            content_lines.append("ПОДПИСИ:")
            content_lines.append("МОЛ: __________________ (подпись) ______ (дата)")
            content_lines.append("Председатель: __________________ (подпись) ______ (дата)")
            content_lines.append("Члены комиссии: __________________ (подпись) ______ (дата)")
            content_lines.append(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
            content_lines.append("=" * 70)

            text_widget.insert('1.0', "\n".join(content_lines))
            text_widget.config(state='disabled')
            tk.Button(print_window, text="🖨️ Печать", command=lambda: self.print_text_widget(text_widget), bg='#3498db', fg='white', width=15).pack(pady=10)

        if can_edit:
            tk.Button(btn_frame, text="💾 Сохранить", command=save_card, bg='#27ae60', fg='white', width=12).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="🖨️ Печать", command=print_card, bg='#3498db', fg='white', width=12).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="✖️ Закрыть", command=dialog.destroy, bg='#95a5a6', fg='white', width=12).pack(side=tk.LEFT, padx=5)

    def show_movements(self):
        self.highlight_menu_button("📋 Движения")
        self.clear_content()

        tk.Label(self.content_frame, text="Движение материалов и основных средств", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        columns = ('item_type', 'item_name', 'operation_type', 'quantity', 'person', 'operation_date', 'created_by')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'item_type': 'Тип', 'item_name': 'Наименование', 'operation_type': 'Операция',
                    'quantity': 'Кол-во', 'person': 'Сотрудник', 'operation_date': 'Дата', 'created_by': 'Кто создал'}

        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=120)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute("""
            SELECT 'Материал' as item_type, m.name, mov.operation_type, mov.quantity, rp.full_name, 
                   mov.operation_date, mov.created_by
            FROM movements mov
            LEFT JOIN materials m ON mov.material_id = m.id
            LEFT JOIN responsible_persons rp ON mov.person_id = rp.id
            WHERE mov.material_id IS NOT NULL
            UNION ALL
            SELECT 'ОС' as item_type, i.name, mov.operation_type, mov.quantity, rp.full_name,
                   mov.operation_date, mov.created_by
            FROM movements mov
            LEFT JOIN inventory_items i ON mov.inventory_item_id = i.id
            LEFT JOIN responsible_persons rp ON mov.person_id = rp.id
            WHERE mov.inventory_item_id IS NOT NULL
            ORDER BY operation_date DESC
            LIMIT 200
        """)
        for row in cur.fetchall():
            op_type = row[2]
            if op_type == 'receipt':
                op_display = '📥 Приход'
            elif op_type == 'issue':
                op_display = '📤 Выдача'
            elif op_type == 'write_off':
                op_display = '❌ Списание'
            elif op_type == 'return':
                op_display = '🔄 Возврат'
            elif op_type == 'move':
                op_display = '📦 Перемещение'
            else:
                op_display = op_type
            values = list(row)
            values[2] = op_display
            tree.insert('', tk.END, values=values)
        cur.close()
        conn.close()

    def show_persons(self):
        self.highlight_menu_button("👥 Сотрудники")
        self.clear_content()

        tk.Label(self.content_frame, text="Материально ответственные лица", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        btn_frame.pack(pady=5)

        if self.role in ['admin', 'storekeeper']:
            tk.Button(btn_frame, text="➕ Добавить сотрудника", command=self.add_person, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="✏️ Редактировать", command=self.edit_person, bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=5)

        if self.role == 'admin':
            tk.Button(btn_frame, text="🗑️ Удалить", command=self.delete_person, bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        search_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        search_frame.pack(pady=5, fill=tk.X)

        tk.Label(search_frame, text="🔍 Поиск:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        self.person_search = tk.Entry(search_frame, width=40)
        self.person_search.pack(side=tk.LEFT, padx=5)
        self.person_search.bind('<KeyRelease>', lambda e: self.load_persons())
        tk.Button(search_frame, text="Очистить", command=self.clear_person_search, bg='#95a5a6', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'full_name', 'position', 'phone', 'email')
        self.person_tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=15)

        headings = {'id': 'ID', 'full_name': 'ФИО', 'position': 'Должность', 'phone': 'Телефон', 'email': 'Email'}
        for col, heading in headings.items():
            self.person_tree.heading(col, text=heading)
            self.person_tree.column(col, width=180)
            self.person_tree.heading(col, command=lambda c=col: self.sort_treeview(self.person_tree, c, False))

        self.person_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.load_persons()

    def load_persons(self):
        for item in self.person_tree.get_children():
            self.person_tree.delete(item)

        search_text = self.person_search.get().strip().lower()

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT id, full_name, position, phone, email FROM responsible_persons ORDER BY full_name")
        for row in cur.fetchall():
            if search_text == "" or search_text in str(row[1]).lower() or (row[2] and search_text in str(row[2]).lower()):
                self.person_tree.insert('', tk.END, values=row)
        cur.close()
        conn.close()

    def clear_person_search(self):
        self.person_search.delete(0, tk.END)
        self.load_persons()

    def add_person(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Добавить сотрудника")
        dialog.geometry("400x400")

        tk.Label(dialog, text="ФИО:").grid(row=0, column=0, padx=10, pady=5, sticky='w')
        name_entry = tk.Entry(dialog, width=35)
        name_entry.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Должность:").grid(row=1, column=0, padx=10, pady=5, sticky='w')
        position_entry = tk.Entry(dialog, width=35)
        position_entry.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Телефон:").grid(row=2, column=0, padx=10, pady=5, sticky='w')
        phone_entry = tk.Entry(dialog, width=35)
        phone_entry.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Email:").grid(row=3, column=0, padx=10, pady=5, sticky='w')
        email_entry = tk.Entry(dialog, width=35)
        email_entry.grid(row=3, column=1, padx=10, pady=5)

        def save():
            full_name = name_entry.get().strip()
            if not full_name:
                messagebox.showerror("Ошибка", "Введите ФИО")
                return

            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute("""
                    INSERT INTO responsible_persons (full_name, position, phone, email)
                    VALUES (%s, %s, %s, %s)
                """, (full_name, position_entry.get().strip(), phone_entry.get().strip(), email_entry.get().strip()))
                conn.commit()
                self.log_action(f"Добавлен сотрудник {full_name}")
                dialog.destroy()
                self.load_persons()
                messagebox.showinfo("Успех", "Сотрудник добавлен")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))
            finally:
                cur.close()
                conn.close()

        tk.Button(dialog, text="Сохранить", command=save, bg='#27ae60', fg='white').grid(row=4, column=0, columnspan=2, pady=20)

    def edit_person(self):
        selected = self.person_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите сотрудника")
            return

        values = self.person_tree.item(selected[0])['values']
        person_id = values[0]
        old_name = values[1]
        old_position = values[2] if values[2] else ''
        old_phone = values[3] if values[3] else ''
        old_email = values[4] if values[4] else ''

        dialog = tk.Toplevel(self.root)
        dialog.title("Редактировать сотрудника")
        dialog.geometry("400x400")

        tk.Label(dialog, text="ФИО:").grid(row=0, column=0, padx=10, pady=5, sticky='w')
        name_entry = tk.Entry(dialog, width=35)
        name_entry.insert(0, old_name)
        name_entry.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Должность:").grid(row=1, column=0, padx=10, pady=5, sticky='w')
        position_entry = tk.Entry(dialog, width=35)
        position_entry.insert(0, old_position)
        position_entry.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Телефон:").grid(row=2, column=0, padx=10, pady=5, sticky='w')
        phone_entry = tk.Entry(dialog, width=35)
        phone_entry.insert(0, old_phone)
        phone_entry.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Email:").grid(row=3, column=0, padx=10, pady=5, sticky='w')
        email_entry = tk.Entry(dialog, width=35)
        email_entry.insert(0, old_email)
        email_entry.grid(row=3, column=1, padx=10, pady=5)

        def save():
            full_name = name_entry.get().strip()
            if not full_name:
                messagebox.showerror("Ошибка", "Введите ФИО")
                return

            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute("""
                    UPDATE responsible_persons 
                    SET full_name=%s, position=%s, phone=%s, email=%s
                    WHERE id=%s
                """, (full_name, position_entry.get().strip(), phone_entry.get().strip(), email_entry.get().strip(), person_id))
                conn.commit()
                self.log_action(f"Изменён сотрудник {old_name} -> {full_name}")
                dialog.destroy()
                self.load_persons()
                messagebox.showinfo("Успех", "Сотрудник обновлён")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))
            finally:
                cur.close()
                conn.close()

        tk.Button(dialog, text="Сохранить", command=save, bg='#f39c12', fg='white').grid(row=4, column=0, columnspan=2, pady=20)

    def delete_person(self):
        selected = self.person_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите сотрудника")
            return

        values = self.person_tree.item(selected[0])['values']
        person_id = values[0]
        person_name = values[1]

        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute("SELECT COUNT(*) FROM movements WHERE person_id = %s", (person_id,))
        movements_count = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM inventory_items WHERE responsible_person_id = %s", (person_id,))
        inventory_count = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM users WHERE person_id = %s", (person_id,))
        user_count = cur.fetchone()[0]

        cur.close()
        conn.close()

        if movements_count > 0 or inventory_count > 0 or user_count > 0:
            message = "Нельзя удалить сотрудника, потому что:\n"
            if movements_count > 0:
                message += f"- есть {movements_count} операций выдачи материалов\n"
            if inventory_count > 0:
                message += f"- на нём числится {inventory_count} единиц основных средств\n"
            if user_count > 0:
                message += f"- у него есть учётная запись в системе\n"
            messagebox.showerror("Ошибка", message)
            return

        if messagebox.askyesno("Подтверждение", f"Удалить сотрудника '{person_name}'?"):
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute("DELETE FROM responsible_persons WHERE id = %s", (person_id,))
                conn.commit()
                self.log_action(f"Удалён сотрудник {person_name}")
                self.load_persons()
                messagebox.showinfo("Успех", "Сотрудник удалён")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))
            finally:
                cur.close()
                conn.close()

    def show_inventory_by_person(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Инвентарная ведомость по МОЛ")
        dialog.geometry("900x600")
        dialog.configure(bg='white')
        # Разворачиваем окно на весь экран (кроссплатформенный вариант)
        try:
            dialog.state('zoomed')  # Windows
        except:
            try:
                dialog.attributes('-zoomed', True)  # Linux
            except:
                dialog.geometry(f"{dialog.winfo_screenwidth()}x{dialog.winfo_screenheight()}+0+0")  # fallback

        main_frame = tk.Frame(dialog, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        tk.Label(main_frame, text="ИНВЕНТАРНАЯ ВЕДОМОСТЬ", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        tk.Label(main_frame, text="Выберите материально ответственное лицо:", font=('Arial', 12), bg='white').pack(pady=5)
        persons = self.get_persons()
        person_combo = ttk.Combobox(main_frame, values=persons, width=50)
        person_combo.pack(pady=5)

        tree_frame = tk.Frame(main_frame, bg='white')
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        columns = ('name', 'inventory_number', 'quantity', 'price', 'total', 'location')
        tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=15)

        headings = {'name': 'Наименование', 'inventory_number': 'Инв.номер', 'quantity': 'Кол-во',
                    'price': 'Цена', 'total': 'Сумма', 'location': 'Место'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150 if col != 'name' else 250)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        total_label = tk.Label(main_frame, text="", font=('Arial', 12, 'bold'), bg='white', fg='#e74c3c')
        total_label.pack(pady=10)

        def load_inventory():
            for item in tree.get_children():
                tree.delete(item)

            person_name = person_combo.get()
            if not person_name:
                return

            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("""
                SELECT i.name, i.inventory_number, i.actual_quantity, i.price, 
                       i.actual_quantity * i.price as total, i.location
                FROM inventory_items i
                LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
                WHERE rp.full_name = %s AND i.status = 'in_use' AND i.actual_quantity > 0
                ORDER BY i.name
            """, (person_name,))
            rows = cur.fetchall()
            cur.close()
            conn.close()

            total_sum = 0
            for row in rows:
                tree.insert('', tk.END, values=row)
                total_sum += row[4] if row[4] else 0

            total_label.config(text=f"ИТОГО: {total_sum:,.2f} руб.")
            return rows, total_sum

        def export_to_excel():
            person_name = person_combo.get()
            if not person_name:
                messagebox.showerror("Ошибка", "Выберите МОЛ")
                return

            rows, total_sum = load_inventory()
            if not rows:
                messagebox.showwarning("Нет данных", "У выбранного МОЛ нет основных средств")
                return

            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel документы", "*.xlsx")],
                initialfile=f"Инвентарная_ведомость_{person_name}.xlsx"
            )

            if not file_path:
                return

            try:
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Инвентарная ведомость"

                ws['A1'] = "ИНВЕНТАРНАЯ ВЕДОМОСТЬ"
                ws.merge_cells('A1:F1')
                ws['A1'].font = Font(size=14, bold=True)
                ws['A1'].alignment = Alignment(horizontal='center')

                ws['A3'] = f"Материально ответственное лицо: {person_name}"
                ws['A4'] = f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}"

                headers = ['Наименование', 'Инв.номер', 'Кол-во', 'Цена', 'Сумма', 'Место']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=6, column=col, value=header)
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    cell.alignment = Alignment(horizontal='center')

                for i, row in enumerate(rows, 7):
                    ws.cell(row=i, column=1, value=row[0])
                    ws.cell(row=i, column=2, value=row[1])
                    ws.cell(row=i, column=3, value=row[2])
                    ws.cell(row=i, column=4, value=row[3])
                    ws.cell(row=i, column=5, value=row[4])
                    ws.cell(row=i, column=6, value=row[5])

                last_row = 7 + len(rows)
                ws.cell(row=last_row, column=4, value="ИТОГО:").font = Font(bold=True)
                ws.cell(row=last_row, column=5, value=total_sum).font = Font(bold=True)

                ws.cell(row=last_row + 2, column=1, value="Материально ответственное лицо: __________________ (подпись) ______ (дата)")

                for col in range(1, 7):
                    ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20

                wb.save(file_path)
                messagebox.showinfo("Успех", f"Ведомость сохранена: {file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        person_combo.bind('<<ComboboxSelected>>', lambda e: load_inventory())
        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="📊 Экспорт в Excel", command=export_to_excel, bg='#1abc9c', fg='white', width=20).pack()

    def show_transfer_act_dialog(self):
        if not self.pending_transfers:
            messagebox.showinfo("Информация", "Нет ожидающих переходов")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Акт перехода основных средств")
        dialog.geometry("900x600")
        dialog.configure(bg='white')

        main_frame = tk.Frame(dialog, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        tk.Label(main_frame, text="АКТ ПЕРЕХОДА ОСНОВНЫХ СРЕДСТВ", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10)

        act_number = f"АП-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        tk.Label(info_frame, text=f"Номер акта: {act_number}", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        tk.Label(info_frame, text=f"Дата: {datetime.now().strftime('%d.%m.%Y')}", font=('Arial', 10), bg='white').grid(row=0, column=1, padx=10, pady=5, sticky='w')

        commission = self.get_commission()
        if commission:
            tk.Label(info_frame, text=f"Комиссия:", font=('Arial', 10, 'bold'), bg='white').grid(row=1, column=0, padx=10, pady=5, sticky='w')
            tk.Label(info_frame, text=f"Председатель: {commission[1]}", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10, pady=2, sticky='w')
            tk.Label(info_frame, text=f"Члены: {commission[2]}", font=('Arial', 10), bg='white', wraplength=400).grid(row=3, column=0, padx=10, pady=2, sticky='w')

        columns = ('name', 'inventory_number', 'from_person', 'to_person', 'from_location', 'to_location')
        tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=10)

        headings = {'name': 'Наименование', 'inventory_number': 'Инв.номер', 'from_person': 'От кого',
                    'to_person': 'Кому', 'from_location': 'Откуда', 'to_location': 'Куда'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150 if col != 'name' else 200)

        tree.pack(fill=tk.BOTH, expand=True, pady=10)

        for transfer in self.pending_transfers:
            item_id, name, location, person_id, _, new_location, new_person_id = transfer
            from_person = self.get_person_name_by_id(person_id) if person_id else 'не указан'
            to_person = self.get_person_name_by_id(new_person_id) if new_person_id else 'не указан'
            tree.insert('', tk.END, values=(name, item_id, from_person, to_person, location or 'не указано', new_location or 'не указано'))

        def confirm_transfer():
            conn = get_db_connection()
            cur = conn.cursor()
            for transfer in self.pending_transfers:
                item_id, name, location, person_id, _, new_location, new_person_id = transfer
                cur.execute("""
                    UPDATE inventory_items SET 
                        location = COALESCE(new_location, location),
                        responsible_person_id = COALESCE(new_person_id, responsible_person_id),
                        transfer_pending = FALSE,
                        new_location = NULL,
                        new_person_id = NULL
                    WHERE id = %s
                """, (item_id,))
                cur.execute("""
                    INSERT INTO movements (inventory_item_id, operation_type, quantity, reason, created_by)
                    VALUES (%s, 'Переход', 1, 'Акт перехода', %s)
                """, (item_id, self.username))
            conn.commit()
            cur.close()
            conn.close()
            self.pending_transfers = []
            self.check_pending_transfers()
            messagebox.showinfo("Успех", "Акт перехода сформирован")
            dialog.destroy()

        def export_to_word():
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word документы", "*.docx")],
                initialfile=f"Акт_перехода_{act_number}.docx"
            )
            if not file_path:
                return

            try:
                doc = Document()
                title = doc.add_heading('АКТ ПЕРЕХОДА ОСНОВНЫХ СРЕДСТВ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Номер акта: {act_number}")
                doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
                if commission:
                    doc.add_paragraph(f"Комиссия: председатель {commission[1]}, члены {commission[2]}")
                doc.add_paragraph()
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Наименование'
                hdr_cells[1].text = 'Инв.номер'
                hdr_cells[2].text = 'От кого'
                hdr_cells[3].text = 'Кому'
                hdr_cells[4].text = 'Новое место'

                for transfer in self.pending_transfers:
                    item_id, name, location, person_id, _, new_location, new_person_id = transfer
                    from_person = self.get_person_name_by_id(person_id) if person_id else 'не указан'
                    to_person = self.get_person_name_by_id(new_person_id) if new_person_id else 'не указан'
                    cells = table.add_row().cells
                    cells[0].text = name
                    cells[1].text = str(item_id)
                    cells[2].text = from_person
                    cells[3].text = to_person
                    cells[4].text = new_location or 'не указано'

                doc.add_paragraph()
                doc.add_paragraph("Подписи:")
                doc.add_paragraph("От过去的МОЛ: _________________________")
                doc.add_paragraph("Новый МОЛ: _________________________")
                doc.add_paragraph("Председатель комиссии: _________________________")
                doc.save(file_path)
                messagebox.showinfo("Успех", f"Акт сохранён: {file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="✅ Подтвердить переход", command=confirm_transfer, bg='#27ae60', fg='white', width=20).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="📎 Экспорт в Word", command=export_to_word, bg='#9b59b6', fg='white', width=20).pack(side=tk.LEFT, padx=10)

    def show_receipt_acts(self):
        self.highlight_menu_button("📄 Акт оприходования")
        self.clear_content()
        tk.Label(self.content_frame, text="Акты оприходования", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        if self.role in ['admin', 'storekeeper']:
            btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
            btn_frame.pack(pady=5)
            tk.Button(btn_frame, text="📄 Создать акт", command=self.create_receipt_act,
                      bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="✏️ Редактировать", command=self.edit_receipt_act,
                      bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="✅ Утвердить", command=self.approve_receipt_act,
                      bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="🗑️ Удалить", command=self.delete_receipt_act,
                      bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📎 Экспорт в Word", command=lambda: self.export_receipt_act_to_word(),
                      bg='#9b59b6', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📊 Экспорт в Excel", command=lambda: self.export_receipt_act_to_excel(),
                      bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'act_number', 'act_date', 'supplier_name', 'receipt_type', 'total_amount', 'status', 'created_by')
        self.receipt_acts_tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=18)

        headings = {'id': 'ID', 'act_number': 'Номер акта', 'act_date': 'Дата',
                    'supplier_name': 'Поставщик', 'receipt_type': 'Тип',
                    'total_amount': 'Сумма', 'status': 'Статус', 'created_by': 'Создал'}
        for col, heading in headings.items():
            self.receipt_acts_tree.heading(col, text=heading)
            self.receipt_acts_tree.column(col, width=100 if col == 'id' else 120)
            self.receipt_acts_tree.heading(col, command=lambda c=col: self.sort_treeview(self.receipt_acts_tree, c, False))

        self.receipt_acts_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.load_receipt_acts()

    def load_receipt_acts(self):
        for item in self.receipt_acts_tree.get_children():
            self.receipt_acts_tree.delete(item)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, act_number, act_date, COALESCE(supplier_name, ''), 
                   receipt_type, total_amount, status, created_by
            FROM receipt_acts
            ORDER BY act_date DESC
        """)
        for row in cur.fetchall():
            self.receipt_acts_tree.insert('', tk.END, values=row)
        cur.close()
        conn.close()

    def delete_receipt_act(self):
        selected = self.receipt_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return
        if messagebox.askyesno("Подтверждение", "Удалить выбранный акт?"):
            act_id = self.receipt_acts_tree.item(selected[0])['values'][0]
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM receipt_act_items WHERE receipt_act_id = %s", (act_id,))
            cur.execute("DELETE FROM receipt_acts WHERE id = %s", (act_id,))
            conn.commit()
            cur.close()
            conn.close()
            self.load_receipt_acts()
            messagebox.showinfo("Успех", "Акт удалён")

    def approve_receipt_act(self):
        selected = self.receipt_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return

        act_id = self.receipt_acts_tree.item(selected[0])['values'][0]
        act_number = self.receipt_acts_tree.item(selected[0])['values'][1]
        current_status = self.receipt_acts_tree.item(selected[0])['values'][6]

        if current_status == 'approved':
            messagebox.showinfo("Информация", "Акт уже утверждён")
            return

        if messagebox.askyesno("Подтверждение",
                               f"Утвердить акт {act_number}? После утверждения позиции будут добавлены на склад."):
            conn = get_db_connection()
            cur = conn.cursor()

            # Получаем позиции акта (включая учётное количество, место и МОЛ)
            cur.execute("""
                SELECT item_type, name, unit_id, quantity, price, accounting_quantity, location, responsible_person
                FROM receipt_act_items WHERE receipt_act_id = %s
            """, (act_id,))
            items = cur.fetchall()

            for item in items:
                item_type, name, unit_id, quantity, price, accounting_quantity, location, responsible_person = item

                if item_type == 'material':
                    # Проверяем, существует ли материал
                    cur.execute("SELECT id, actual_quantity, accounting_quantity FROM materials WHERE name = %s",
                                (name,))
                    existing = cur.fetchone()
                    if existing:
                        # Обновляем существующий материал (просто добавляем количество, без корректировки)
                        new_actual = existing[1] + quantity
                        new_accounting = existing[2] + (accounting_quantity if accounting_quantity else quantity)
                        cur.execute("""
                            UPDATE materials SET actual_quantity = %s, accounting_quantity = %s, location = %s
                            WHERE id = %s
                        """, (new_actual, new_accounting,
                              location if location else existing[3] if len(existing) > 3 else None, existing[0]))
                    else:
                        # Создаём новый материал
                        cur.execute("""
                            INSERT INTO materials (name, unit_id, price, actual_quantity, accounting_quantity, location, status)
                            VALUES (%s, %s, %s, %s, %s, %s, 'active')
                        """, (name, unit_id, price, quantity, accounting_quantity if accounting_quantity else quantity,
                              location))

                elif item_type == 'inventory':
                    # Для ОС: извлекаем инвентарный номер из имени (если он там был)
                    import re
                    inv_match = re.search(r'\(инв\.№([^)]+)\)', name)
                    if inv_match:
                        inventory_number = inv_match.group(1)
                        clean_name = re.sub(r'\s*\(инв\.№[^)]+\)', '', name).strip()
                    else:
                        inventory_number = f"ИНВ-{datetime.now().strftime('%Y%m%d%H%M%S')}{act_id}"
                        clean_name = name

                    # Получаем ID ответственного лица
                    person_id = None
                    if responsible_person:
                        cur.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (responsible_person,))
                        pers = cur.fetchone()
                        if pers:
                            person_id = pers[0]
                        else:
                            cur.execute("INSERT INTO responsible_persons (full_name) VALUES (%s) RETURNING id",
                                        (responsible_person,))
                            person_id = cur.fetchone()[0]

                    cur.execute("""
                        INSERT INTO inventory_items (inventory_number, name, unit_id, price, actual_quantity, accounting_quantity,
                                                     location, responsible_person_id, commissioning_date, status)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, CURRENT_DATE, 'in_use')
                    """, (inventory_number, clean_name, unit_id, price, quantity,
                          accounting_quantity if accounting_quantity else quantity, location, person_id))

            # Обновляем статус акта
            cur.execute("UPDATE receipt_acts SET status = 'approved' WHERE id = %s", (act_id,))
            conn.commit()
            cur.close()
            conn.close()

            # Обновляем списки
            self.load_receipt_acts()

            if hasattr(self, 'material_tree') and self.material_tree:
                try:
                    self.load_materials_data()
                except:
                    pass
            if hasattr(self, 'inventory_tree') and self.inventory_tree:
                try:
                    self.load_inventory_data()
                except:
                    pass

            messagebox.showinfo("Успех", f"Акт {act_number} утверждён. Позиции добавлены на склад.")
            # Проверяем наличие расхождений и показываем уведомление
            self.show_discrepancy_notification()

    def edit_receipt_act(self):
        selected = self.receipt_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return

        act_id = self.receipt_acts_tree.item(selected[0])['values'][0]
        current_status = self.receipt_acts_tree.item(selected[0])['values'][6]

        if current_status == 'approved':
            messagebox.showwarning("Ошибка", "Нельзя редактировать утверждённый акт")
            return

        self.edit_receipt_act_dialog(act_id)

    def edit_receipt_act_dialog(self, act_id):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, act_number, act_date, supplier_name, receipt_type, responsible_person_id, total_amount, status, created_by
            FROM receipt_acts WHERE id = %s
        """, (act_id,))
        act = cur.fetchone()

        cur.execute("""
            SELECT id, item_type, name, unit_id, quantity, price, total
            FROM receipt_act_items WHERE receipt_act_id = %s
        """, (act_id,))
        items = cur.fetchall()
        cur.close()
        conn.close()

        if not act:
            messagebox.showerror("Ошибка", "Акт не найден")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Редактирование акта оприходования - {act[1]}")
        dialog.geometry("900x700")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.focus_set()

        main_frame = tk.Frame(dialog, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        tk.Label(main_frame, text="АКТ ОПРИХОДОВАНИЯ (РЕДАКТИРОВАНИЕ)", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10)

        tk.Label(info_frame, text="Номер акта:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        tk.Label(info_frame, text=act[1], font=('Arial', 10, 'bold'), bg='white', fg='#27ae60').grid(row=0, column=1, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Дата акта:", font=('Arial', 10), bg='white').grid(row=0, column=2, padx=10, pady=5, sticky='w')
        date_entry = tk.Entry(info_frame, width=15)
        date_entry.insert(0, act[2].strftime('%Y-%m-%d') if act[2] else datetime.now().strftime('%Y-%m-%d'))
        date_entry.grid(row=0, column=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Поставщик:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10, pady=5, sticky='w')
        suppliers = self.get_suppliers()
        supplier_combo = ttk.Combobox(info_frame, values=suppliers, width=40)
        supplier_combo.set(act[3] if act[3] else '')
        supplier_combo.grid(row=1, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Тип поступления:", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10, pady=5, sticky='w')
        receipt_type_combo = ttk.Combobox(info_frame, values=['покупка', 'безвозмездно'], width=40)
        receipt_type_combo.set('покупка' if act[4] == 'purchase' else 'безвозмездно')
        receipt_type_combo.grid(row=2, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Ответственное лицо:", font=('Arial', 10), bg='white').grid(row=3, column=0, padx=10, pady=5, sticky='w')
        persons = self.get_persons()
        person_combo = ttk.Combobox(info_frame, values=persons, width=40)
        person_name = self.get_person_name_by_id(act[5]) if act[5] else ''
        person_combo.set(person_name)
        person_combo.grid(row=3, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        items_frame = tk.LabelFrame(main_frame, text="Позиции оприходования", font=('Arial', 12, 'bold'), bg='white')
        items_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        columns = ('id', 'item_type', 'name', 'unit', 'quantity', 'price', 'total')
        items_tree = ttk.Treeview(items_frame, columns=columns, show='headings', height=10)

        headings = {'id': 'ID', 'item_type': 'Тип', 'name': 'Наименование', 'unit': 'Ед.изм',
                    'quantity': 'Кол-во', 'price': 'Цена', 'total': 'Сумма'}
        for col, heading in headings.items():
            items_tree.heading(col, text=heading)
            items_tree.column(col, width=80 if col == 'id' else 150 if col != 'name' else 200)

        items_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        for item in items:
            unit_name = self.get_unit_name_by_id(item[3]) if item[3] else ''
            items_tree.insert('', tk.END, values=(item[0], item[1], item[2], unit_name, item[4], item[5], item[6]))

        btn_items_frame = tk.Frame(items_frame, bg='white')
        btn_items_frame.pack(pady=5)

        tk.Button(btn_items_frame, text="➕ Добавить материал",
                  command=lambda: self.add_receipt_item_dialog(dialog, items_tree, 'material', self.get_units()),
                  bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_items_frame, text="➕ Добавить ОС",
                  command=lambda: self.add_receipt_item_dialog(dialog, items_tree, 'inventory', self.get_units()),
                  bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_items_frame, text="🗑️ Удалить",
                  command=lambda: self.delete_selected_tree_item(items_tree),
                  bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        total_frame = tk.Frame(main_frame, bg='white')
        total_frame.pack(fill=tk.X, pady=10)
        tk.Label(total_frame, text="ИТОГО:", font=('Arial', 12, 'bold'), bg='white', fg='#e74c3c').pack(side=tk.RIGHT, padx=10)
        total_label = tk.Label(total_frame, text=f"{act[6]:,.2f} руб.", font=('Arial', 12, 'bold'), bg='white', fg='#e74c3c')
        total_label.pack(side=tk.RIGHT, padx=10)

        # !!! ФУНКЦИЯ update_total ОПРЕДЕЛЕНА ЗДЕСЬ (ДО её вызова) !!!
        def update_total():
            total = 0
            for item in items_tree.get_children():
                values = items_tree.item(item)['values']
                if len(values) >= 7 and values[6]:
                    try:
                        total += float(values[6])
                    except:
                        pass
                elif len(values) >= 6 and values[5]:
                    try:
                        total += float(values[5])
                    except:
                        pass
            total_label.config(text=f"{total:,.2f} руб.")
            return total

        def save_act():
            try:
                act_date = date_entry.get().strip()
                supplier_name = supplier_combo.get().strip()
                receipt_type = 'purchase' if receipt_type_combo.get() == 'покупка' else 'free'
                person_name = person_combo.get().strip()

                if not supplier_name:
                    messagebox.showerror("Ошибка", "Введите поставщика")
                    return

                if items_tree.get_children() == ():
                    messagebox.showerror("Ошибка", "Добавьте хотя бы одну позицию")
                    return

                total_sum = update_total()  # <-- вызов функции ПОСЛЕ её определения

                conn2 = get_db_connection()
                cur2 = conn2.cursor()

                cur2.execute("SELECT id FROM suppliers WHERE name = %s", (supplier_name,))
                sup = cur2.fetchone()
                if sup:
                    supplier_id = sup[0]
                else:
                    cur2.execute("INSERT INTO suppliers (name) VALUES (%s) RETURNING id", (supplier_name,))
                    supplier_id = cur2.fetchone()[0]

                person_id = None
                if person_name:
                    cur2.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (person_name,))
                    pers = cur2.fetchone()
                    if pers:
                        person_id = pers[0]
                    else:
                        cur2.execute("INSERT INTO responsible_persons (full_name) VALUES (%s) RETURNING id", (person_name,))
                        person_id = cur2.fetchone()[0]

                cur2.execute("""
                    UPDATE receipt_acts SET 
                        act_date=%s, supplier_id=%s, supplier_name=%s, receipt_type=%s, 
                        responsible_person_id=%s, total_amount=%s
                    WHERE id=%s
                """, (act_date, supplier_id, supplier_name, receipt_type, person_id, total_sum, act_id))

                cur2.execute("DELETE FROM receipt_act_items WHERE receipt_act_id = %s", (act_id,))

                for item in items_tree.get_children():
                    values = items_tree.item(item)['values']
                    item_type = values[1]
                    name = values[2]
                    unit = values[3]
                    quantity = float(values[4]) if values[4] else 0
                    price = float(values[5]) if values[5] else 0
                    total = quantity * price

                    unit_id = None
                    if unit:
                        cur2.execute("SELECT id FROM units WHERE short_name = %s OR full_name = %s", (unit, unit))
                        u = cur2.fetchone()
                        if u:
                            unit_id = u[0]

                    cur2.execute("""
                        INSERT INTO receipt_act_items (receipt_act_id, item_type, name, unit_id, quantity, price, total)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """, (act_id, item_type, name, unit_id, quantity, price, total))

                conn2.commit()
                cur2.close()
                conn2.close()

                messagebox.showinfo("Успех", f"Акт {act[1]} обновлён")
                dialog.destroy()
                self.load_receipt_acts()
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="💾 Сохранить изменения", command=save_act, bg='#27ae60', fg='white', width=20).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="✖️ Отмена", command=dialog.destroy, bg='#95a5a6', fg='white', width=15).pack(side=tk.LEFT, padx=10)

    def export_receipt_act_to_word(self):
        selected = self.receipt_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return
        if not DOCX_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека python-docx не установлена")
            return
        act_id = self.receipt_acts_tree.item(selected[0])['values'][0]
        self.export_receipt_act_to_word_by_id(act_id)

    def export_receipt_act_to_word_by_id(self, act_id, file_path=None):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT ra.act_number, ra.act_date, ra.supplier_name, ra.receipt_type, ra.total_amount,
                   ra.created_by, ra.created_at, ri.name, ri.quantity, ri.price, ri.total, u.short_name
            FROM receipt_acts ra
            LEFT JOIN receipt_act_items ri ON ra.id = ri.receipt_act_id
            LEFT JOIN units u ON ri.unit_id = u.id
            WHERE ra.id = %s
        """, (act_id,))
        rows = cur.fetchall()
        cur.close()
        conn.close()

        if not rows:
            messagebox.showerror("Ошибка", "Акт не найден")
            return

        if file_path is None:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word документы", "*.docx")],
                initialfile=f"Акт_оприходования_{rows[0][0]}.docx"
            )
            if not file_path:
                return

        try:
            doc = Document()
            title = doc.add_heading('АКТ ОПРИХОДОВАНИЯ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Номер акта: {rows[0][0]}")
            doc.add_paragraph(f"Дата: {rows[0][1]}")
            doc.add_paragraph(f"Поставщик: {rows[0][2]}")
            doc.add_paragraph(f"Тип поступления: {rows[0][3]}")
            doc.add_paragraph(f"Создал: {rows[0][5]}")
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Наименование'
            hdr_cells[1].text = 'Ед.изм'
            hdr_cells[2].text = 'Количество'
            hdr_cells[3].text = 'Цена'
            hdr_cells[4].text = 'Сумма'
            for row in rows:
                cells = table.add_row().cells
                cells[0].text = row[7] if row[7] else ''
                cells[1].text = row[11] if row[11] else ''
                cells[2].text = str(row[8]) if row[8] else ''
                cells[3].text = f"{row[9]:.2f}" if row[9] else ''
                cells[4].text = f"{row[10]:.2f}" if row[10] else ''
            doc.add_paragraph()
            doc.add_paragraph(f"ИТОГО: {rows[0][4]:,.2f} руб." if rows[0][4] else "ИТОГО: 0.00 руб.")
            doc.add_paragraph()
            doc.add_paragraph("Подписи: _________________________")
            doc.save(file_path)
            messagebox.showinfo("Успех", f"Акт сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def export_receipt_act_to_excel(self):
        selected = self.receipt_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return
        if not EXCEL_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека openpyxl не установлена")
            return
        act_id = self.receipt_acts_tree.item(selected[0])['values'][0]
        self.export_receipt_act_to_excel_by_id(act_id)

    def export_receipt_act_to_excel_by_id(self, act_id, file_path=None):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT ra.act_number, ra.act_date, ra.supplier_name, ra.receipt_type, ra.total_amount,
                   ra.created_by, ra.created_at, ri.name, ri.quantity, ri.price, ri.total, u.short_name
            FROM receipt_acts ra
            LEFT JOIN receipt_act_items ri ON ra.id = ri.receipt_act_id
            LEFT JOIN units u ON ri.unit_id = u.id
            WHERE ra.id = %s
        """, (act_id,))
        rows = cur.fetchall()
        cur.close()
        conn.close()

        if not rows:
            messagebox.showerror("Ошибка", "Акт не найден")
            return

        if file_path is None:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel документы", "*.xlsx")],
                initialfile=f"Акт_оприходования_{rows[0][0]}.xlsx"
            )
            if not file_path:
                return

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Акт оприходования"
            ws['A1'] = "АКТ ОПРИХОДОВАНИЯ"
            ws.merge_cells('A1:E1')
            ws['A1'].font = Font(size=16, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')
            ws['A3'] = f"Номер акта: {rows[0][0]}"
            ws['A4'] = f"Дата: {rows[0][1]}"
            ws['A5'] = f"Поставщик: {rows[0][2]}"
            ws['A6'] = f"Тип поступления: {rows[0][3]}"
            ws['A7'] = f"Создал: {rows[0][5]}"
            headers = ['Наименование', 'Ед.изм', 'Количество', 'Цена', 'Сумма']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=9, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')
            for i, row in enumerate(rows, 10):
                ws.cell(row=i, column=1, value=row[7] if row[7] else '')
                ws.cell(row=i, column=2, value=row[11] if row[11] else '')
                ws.cell(row=i, column=3, value=row[8] if row[8] else 0)
                ws.cell(row=i, column=4, value=row[9] if row[9] else 0)
                ws.cell(row=i, column=5, value=row[10] if row[10] else 0)
            last_row = 10 + len(rows)
            ws.cell(row=last_row, column=4, value="ИТОГО:").font = Font(bold=True)
            ws.cell(row=last_row, column=5, value=rows[0][4] if rows[0][4] else 0).font = Font(bold=True)
            for col in range(1, 6):
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20
            wb.save(file_path)
            messagebox.showinfo("Успех", f"Акт сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def show_write_off_acts(self):
        self.highlight_menu_button("📑 Акт списания")
        self.clear_content()
        tk.Label(self.content_frame, text="Акты списания", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        if self.role in ['admin', 'storekeeper']:
            btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
            btn_frame.pack(pady=5)
            tk.Button(btn_frame, text="📄 Создать акт", command=self.create_write_off_act,
                      bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="✏️ Редактировать", command=self.edit_write_off_act,
                      bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="✅ Утвердить", command=self.approve_write_off_act,
                      bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📎 Экспорт в Word", command=self.export_write_off_act_to_word,
                      bg='#9b59b6', fg='white').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="📊 Экспорт в Excel", command=self.export_write_off_act_to_excel,
                      bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'act_number', 'act_date', 'write_off_type', 'total_amount', 'status', 'created_by')
        self.write_off_acts_tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=18)

        headings = {'id': 'ID', 'act_number': 'Номер акта', 'act_date': 'Дата', 'write_off_type': 'Тип',
                    'total_amount': 'Сумма', 'status': 'Статус', 'created_by': 'Создал'}
        for col, heading in headings.items():
            self.write_off_acts_tree.heading(col, text=heading)
            self.write_off_acts_tree.column(col, width=120)
            self.write_off_acts_tree.heading(col, command=lambda c=col: self.sort_treeview(self.write_off_acts_tree, c, False))

        self.write_off_acts_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.load_write_off_acts()

    def load_write_off_acts(self):
        for item in self.write_off_acts_tree.get_children():
            self.write_off_acts_tree.delete(item)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, act_number, act_date, write_off_type, total_amount, status, created_by
            FROM write_off_acts
            ORDER BY act_date DESC
        """)
        for row in cur.fetchall():
            self.write_off_acts_tree.insert('', tk.END, values=row)
        cur.close()
        conn.close()

    def approve_write_off_act(self):
        selected = self.write_off_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return

        act_id = self.write_off_acts_tree.item(selected[0])['values'][0]
        act_number = self.write_off_acts_tree.item(selected[0])['values'][1]
        write_off_type = self.write_off_acts_tree.item(selected[0])['values'][3]
        current_status = self.write_off_acts_tree.item(selected[0])['values'][5]

        if current_status == 'approved':
            messagebox.showinfo("Информация", "Акт уже утверждён")
            return

        if messagebox.askyesno("Подтверждение",
                               f"Утвердить акт списания {act_number}? После утверждения позиции будут списаны со склада."):
            conn = get_db_connection()
            cur = conn.cursor()

            cur.execute("""
                SELECT item_type, name, quantity
                FROM write_off_act_items WHERE write_off_act_id = %s
            """, (act_id,))
            items = cur.fetchall()

            for item in items:
                item_type, name, quantity = item

                if write_off_type == 'материалы' or item_type == 'материалы':
                    cur.execute("SELECT id, actual_quantity FROM materials WHERE name = %s", (name,))
                    existing = cur.fetchone()
                    if existing:
                        new_quantity = max(0, existing[1] - quantity)
                        cur.execute("""
                            UPDATE materials SET actual_quantity = %s
                            WHERE id = %s
                        """, (new_quantity, existing[0]))

                elif write_off_type == 'основные_средства' or item_type == 'основные_средства':
                    cur.execute("""
                        UPDATE inventory_items SET status = 'written_off', actual_quantity = 0
                        WHERE name = %s AND status = 'in_use'
                    """, (name,))

            cur.execute("UPDATE write_off_acts SET status = 'approved' WHERE id = %s", (act_id,))
            conn.commit()
            cur.close()
            conn.close()

            self.load_write_off_acts()

            if hasattr(self, 'material_tree') and self.material_tree:
                try:
                    self.load_materials_data()
                except:
                    pass
            if hasattr(self, 'inventory_tree') and self.inventory_tree:
                try:
                    self.load_inventory_data()
                except:
                    pass

            messagebox.showinfo("Успех", f"Акт списания {act_number} утверждён. Позиции списаны.")

    def edit_write_off_act(self):
        selected = self.write_off_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return

        act_id = self.write_off_acts_tree.item(selected[0])['values'][0]
        current_status = self.write_off_acts_tree.item(selected[0])['values'][5]

        if current_status == 'approved':
            messagebox.showwarning("Ошибка", "Нельзя редактировать утверждённый акт")
            return

        self.edit_write_off_act_dialog(act_id)

    def edit_write_off_act_dialog(self, act_id):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, act_number, act_date, write_off_type, total_amount, status, reason, created_by, commission_id
            FROM write_off_acts WHERE id = %s
        """, (act_id,))
        act = cur.fetchone()

        cur.execute("""
            SELECT id, item_type, name, unit_id, quantity, price, total, reason
            FROM write_off_act_items WHERE write_off_act_id = %s
        """, (act_id,))
        items = cur.fetchall()
        cur.close()
        conn.close()

        if not act:
            messagebox.showerror("Ошибка", "Акт не найден")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Редактирование акта списания - {act[1]}")
        dialog.geometry("800x650")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.focus_set()

        main_frame = tk.Frame(dialog, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        tk.Label(main_frame, text="АКТ СПИСАНИЯ (РЕДАКТИРОВАНИЕ)", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10)

        tk.Label(info_frame, text="Номер акта:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        tk.Label(info_frame, text=act[1], font=('Arial', 10, 'bold'), bg='white', fg='#27ae60').grid(row=0, column=1, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Дата акта:", font=('Arial', 10), bg='white').grid(row=0, column=2, padx=10, pady=5, sticky='w')
        date_entry = tk.Entry(info_frame, width=15)
        date_entry.insert(0, act[2].strftime('%Y-%m-%d') if act[2] else datetime.now().strftime('%Y-%m-%d'))
        date_entry.grid(row=0, column=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Тип списания:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10, pady=5, sticky='w')
        type_combo = ttk.Combobox(info_frame, values=['материалы', 'основные_средства'], width=40)
        type_combo.set(act[3])
        type_combo.grid(row=1, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Причина списания:", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10, pady=5, sticky='w')
        reason_entry = tk.Text(info_frame, height=3, width=40)
        reason_entry.insert('1.0', act[6] if act[6] else '')
        reason_entry.grid(row=2, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Комиссия:", font=('Arial', 10), bg='white').grid(row=3, column=0, padx=10, pady=5, sticky='w')
        commission = self.get_commission()
        commission_label = tk.Label(info_frame, text=commission[1] if commission else "Не назначена", font=('Arial', 10), bg='white')
        commission_label.grid(row=3, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        items_frame = tk.LabelFrame(main_frame, text="Позиции списания", font=('Arial', 12, 'bold'), bg='white')
        items_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        columns = ('id', 'name', 'unit', 'quantity', 'price', 'total')
        items_tree = ttk.Treeview(items_frame, columns=columns, show='headings', height=10)

        headings = {'id': 'ID', 'name': 'Наименование', 'unit': 'Ед.изм', 'quantity': 'Кол-во', 'price': 'Цена', 'total': 'Сумма'}
        for col, heading in headings.items():
            items_tree.heading(col, text=heading)
            items_tree.column(col, width=100 if col == 'id' else 150 if col != 'name' else 200)

        items_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        for item in items:
            unit_name = self.get_unit_name_by_id(item[3]) if item[3] else ''
            items_tree.insert('', tk.END, values=(item[0], item[2], unit_name, item[4], item[5], item[6]))

        btn_items_frame = tk.Frame(items_frame, bg='white')
        btn_items_frame.pack(pady=5)

        tk.Button(btn_items_frame, text="➕ Добавить",
                  command=lambda: self.add_write_off_item_dialog(dialog, items_tree, type_combo.get(), self.get_units()),
                  bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_items_frame, text="🗑️ Удалить",
                  command=lambda: self.delete_selected_tree_item(items_tree),
                  bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        total_frame = tk.Frame(main_frame, bg='white')
        total_frame.pack(fill=tk.X, pady=10)
        tk.Label(total_frame, text="ИТОГО:", font=('Arial', 12, 'bold'), bg='white', fg='#e74c3c').pack(side=tk.RIGHT, padx=10)
        total_label = tk.Label(total_frame, text=f"{act[4]:,.2f} руб.", font=('Arial', 12, 'bold'), bg='white', fg='#e74c3c')
        total_label.pack(side=tk.RIGHT, padx=10)

        # !!! ФУНКЦИЯ update_total ОПРЕДЕЛЕНА ЗДЕСЬ (ДО её вызова) !!!
        def update_total():
            total = 0
            for item in items_tree.get_children():
                values = items_tree.item(item)['values']
                if len(values) >= 6 and values[5]:
                    try:
                        total += float(values[5])
                    except:
                        pass
                elif len(values) >= 5 and values[4]:
                    try:
                        total += float(values[4])
                    except:
                        pass
            total_label.config(text=f"{total:,.2f} руб.")
            return total

        def save_act():
            try:
                act_date = date_entry.get().strip()
                write_off_type = type_combo.get()
                reason = reason_entry.get('1.0', tk.END).strip()

                if items_tree.get_children() == ():
                    messagebox.showerror("Ошибка", "Добавьте хотя бы одну позицию")
                    return

                total_sum = update_total()  # <-- вызов функции ПОСЛЕ её определения

                conn2 = get_db_connection()
                cur2 = conn2.cursor()

                commission_id = commission[0] if commission else None

                cur2.execute("""
                    UPDATE write_off_acts SET 
                        act_date=%s, write_off_type=%s, total_amount=%s, reason=%s, commission_id=%s
                    WHERE id=%s
                """, (act_date, write_off_type, total_sum, reason, commission_id, act_id))

                cur2.execute("DELETE FROM write_off_act_items WHERE write_off_act_id = %s", (act_id,))

                for item in items_tree.get_children():
                    values = items_tree.item(item)['values']
                    name = values[1]
                    unit = values[2] if len(values) > 2 else ''
                    quantity = float(values[3]) if len(values) > 3 and values[3] else 0
                    price = float(values[4]) if len(values) > 4 and values[4] else 0
                    total = quantity * price

                    unit_id = None
                    if unit:
                        cur2.execute("SELECT id FROM units WHERE short_name = %s OR full_name = %s", (unit, unit))
                        u = cur2.fetchone()
                        if u:
                            unit_id = u[0]

                    cur2.execute("""
                        INSERT INTO write_off_act_items (write_off_act_id, item_type, name, unit_id, quantity, price, total, reason)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """, (act_id, write_off_type, name, unit_id, quantity, price, total, reason))

                conn2.commit()
                cur2.close()
                conn2.close()

                messagebox.showinfo("Успех", f"Акт {act[1]} обновлён")
                dialog.destroy()
                self.load_write_off_acts()
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="💾 Сохранить изменения", command=save_act, bg='#27ae60', fg='white', width=20).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="✖️ Отмена", command=dialog.destroy, bg='#95a5a6', fg='white', width=15).pack(side=tk.LEFT, padx=10)

    def export_write_off_act_to_word(self):
        selected = self.write_off_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return
        if not DOCX_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека python-docx не установлена")
            return

        values = self.write_off_acts_tree.item(selected[0])['values']
        act_number = values[1]
        act_date = values[2]
        write_off_type = values[3]
        total_amount = values[4]
        created_by = values[6]

        try:
            total_amount_float = float(total_amount) if total_amount else 0
        except (ValueError, TypeError):
            total_amount_float = 0

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word документы", "*.docx")],
            initialfile=f"Акт_списания_{act_number}.docx"
        )
        if not file_path:
            return

        try:
            doc = Document()
            title = doc.add_heading('АКТ СПИСАНИЯ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Номер акта: {act_number}")
            doc.add_paragraph(f"Дата: {act_date}")
            doc.add_paragraph(f"Тип списания: {write_off_type}")
            doc.add_paragraph(f"Создал: {created_by}")
            doc.add_paragraph()
            doc.add_paragraph(f"ИТОГО: {total_amount_float:,.2f} руб." if total_amount_float else "ИТОГО: 0.00 руб.")
            doc.add_paragraph()
            doc.add_paragraph("Председатель комиссии: _________________________")
            doc.add_paragraph("Члены комиссии: _________________________")
            doc.save(file_path)
            messagebox.showinfo("Успех", f"Акт сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта: {str(e)}")

    def export_write_off_act_to_excel(self):
        selected = self.write_off_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return
        if not EXCEL_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека openpyxl не установлена")
            return

        values = self.write_off_acts_tree.item(selected[0])['values']
        act_number = values[1]
        act_date = values[2]
        write_off_type = values[3]
        total_amount = values[4]
        created_by = values[6]

        try:
            total_amount_float = float(total_amount) if total_amount else 0
        except (ValueError, TypeError):
            total_amount_float = 0

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel документы", "*.xlsx")],
            initialfile=f"Акт_списания_{act_number}.xlsx"
        )
        if not file_path:
            return

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Акт списания"
            ws['A1'] = "АКТ СПИСАНИЯ"
            ws.merge_cells('A1:D1')
            ws['A1'].font = Font(size=16, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')
            ws['A3'] = f"Номер акта: {act_number}"
            ws['A4'] = f"Дата: {act_date}"
            ws['A5'] = f"Тип списания: {write_off_type}"
            ws['A6'] = f"Создал: {created_by}"
            ws['A8'] = f"ИТОГО: {total_amount_float:,.2f} руб." if total_amount_float else "ИТОГО: 0.00 руб."
            for col in range(1, 5):
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20
            wb.save(file_path)
            messagebox.showinfo("Успех", f"Акт сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта: {str(e)}")

    def show_discrepancy_acts(self):
        self.highlight_menu_button("⚠️ Акт о расхождениях")
        self.clear_content()

        tk.Label(self.content_frame, text="Акты о расхождениях", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        if self.role in ['admin', 'storekeeper']:
            btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
            btn_frame.pack(pady=5)
            tk.Button(btn_frame, text="📄 Создать акт", command=self.create_discrepancy_act,
                      bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        tk.Label(self.content_frame, text="Текущие выявленные расхождения", font=('Arial', 14, 'bold'),
                 bg='#ecf0f1', fg='#e74c3c').pack(pady=5)

        columns = ('item_type', 'item_name', 'inventory_number', 'location', 'responsible_person',
                   'surplus_quantity', 'surplus_sum', 'shortage_quantity', 'shortage_sum')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=10)

        headings = {'item_type': 'Тип', 'item_name': 'Наименование', 'inventory_number': 'Инв.номер',
                    'location': 'Место', 'responsible_person': 'МОЛ',
                    'surplus_quantity': 'Излишки (кол-во)', 'surplus_sum': 'Излишки (сумма)',
                    'shortage_quantity': 'Недостача (кол-во)', 'shortage_sum': 'Недостача (сумма)'}

        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=120 if col not in ['item_name', 'responsible_person'] else 150)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        conn = get_db_connection()
        cur = conn.cursor()

        # Исправленный запрос - получаем location и ответственное лицо из соответствующих таблиц
        cur.execute("""
            SELECT 
                nv.item_type, 
                nv.item_name,
                CASE 
                    WHEN nv.item_type = 'inventory' THEN COALESCE(i.inventory_number, '')
                    ELSE ''
                END as inventory_number,
                CASE 
                    WHEN nv.item_type = 'material' THEN COALESCE(m.location, '')
                    WHEN nv.item_type = 'inventory' THEN COALESCE(i.location, '')
                    ELSE ''
                END as location,
                CASE 
                    WHEN nv.item_type = 'inventory' THEN COALESCE(rp.full_name, '')
                    ELSE ''
                END as responsible_person,
                nv.surplus_quantity, 
                nv.surplus_sum, 
                nv.shortage_quantity, 
                nv.shortage_sum
            FROM notifications_view nv
            LEFT JOIN inventory_items i ON i.name = nv.item_name AND nv.item_type = 'inventory'
            LEFT JOIN materials m ON m.name = nv.item_name AND nv.item_type = 'material'
            LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
            WHERE nv.surplus_quantity != 0 OR nv.shortage_quantity != 0
            ORDER BY nv.item_type, nv.item_name
        """)
        for row in cur.fetchall():
            tree.insert('', tk.END, values=row)
        cur.close()
        conn.close()

        tk.Label(self.content_frame, text="Сохранённые акты о расхождениях", font=('Arial', 14, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        columns_acts = ('id', 'act_number', 'act_date', 'total_surplus_sum', 'total_shortage_sum', 'status',
                        'created_by')
        self.discrepancy_acts_tree = ttk.Treeview(self.content_frame, columns=columns_acts, show='headings', height=8)

        headings_acts = {'id': 'ID', 'act_number': 'Номер акта', 'act_date': 'Дата',
                         'total_surplus_sum': 'Сумма излишков', 'total_shortage_sum': 'Сумма недостачи',
                         'status': 'Статус', 'created_by': 'Создал'}
        for col, heading in headings_acts.items():
            self.discrepancy_acts_tree.heading(col, text=heading)
            self.discrepancy_acts_tree.column(col, width=120)
            self.discrepancy_acts_tree.heading(col,
                                               command=lambda c=col: self.sort_treeview(self.discrepancy_acts_tree, c,
                                                                                        False))

        self.discrepancy_acts_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        acts_btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        acts_btn_frame.pack(pady=5)
        tk.Button(acts_btn_frame, text="📄 Просмотреть выбранный акт", command=self.view_discrepancy_act_details,
                  bg='#3498db', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(acts_btn_frame, text="✅ Утвердить акт", command=self.approve_discrepancy_act,
                  bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(acts_btn_frame, text="📎 Экспорт в Excel", command=self.export_discrepancy_act,
                  bg='#9b59b6', fg='white').pack(side=tk.LEFT, padx=5)

        self.load_discrepancy_acts_list()

    def load_discrepancy_acts_list(self):
        for item in self.discrepancy_acts_tree.get_children():
            self.discrepancy_acts_tree.delete(item)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, act_number, act_date, total_surplus_sum, total_shortage_sum, status, created_by
            FROM discrepancy_acts
            ORDER BY act_date DESC
        """)
        for row in cur.fetchall():
            self.discrepancy_acts_tree.insert('', tk.END, values=row)
        cur.close()
        conn.close()

    def approve_discrepancy_act(self):
        selected = self.discrepancy_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return

        act_id = self.discrepancy_acts_tree.item(selected[0])['values'][0]
        act_number = self.discrepancy_acts_tree.item(selected[0])['values'][1]
        current_status = self.discrepancy_acts_tree.item(selected[0])['values'][5]

        if current_status == 'approved':
            messagebox.showinfo("Информация", "Акт уже утверждён")
            return

        # Получаем позиции акта о расхождениях
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT dai.id, dai.item_type, dai.surplus_quantity, dai.shortage_quantity,
                   nv.item_name
            FROM discrepancy_act_items dai
            LEFT JOIN notifications_view nv ON nv.item_type = dai.item_type
            WHERE dai.discrepancy_act_id = %s
        """, (act_id,))
        items = cur.fetchall()
        cur.close()
        conn.close()

        if not items:
            messagebox.showerror("Ошибка", "Акт не содержит позиций")
            return

        # Формируем подробное сообщение о корректировке
        msg = "⚠️ ВНИМАНИЕ! При утверждении этого акта будут скорректированы остатки:\n\n"
        msg += "=" * 60 + "\n"
        total_surplus = 0
        total_shortage = 0
        corrections = []

        for item in items:
            item_id, item_type, surplus_qty, shortage_qty, item_name = item
            if surplus_qty and surplus_qty > 0:
                msg += f"📈 {item_name} ({item_type}): ИЗЛИШКИ +{surplus_qty} шт.\n"
                total_surplus += surplus_qty
                corrections.append({
                    'item_id': item_id,
                    'item_type': item_type,
                    'item_name': item_name,
                    'surplus': surplus_qty,
                    'shortage': 0
                })
            if shortage_qty and shortage_qty > 0:
                msg += f"📉 {item_name} ({item_type}): НЕДОСТАЧА -{shortage_qty} шт.\n"
                total_shortage += shortage_qty
                corrections.append({
                    'item_id': item_id,
                    'item_type': item_type,
                    'item_name': item_name,
                    'surplus': 0,
                    'shortage': shortage_qty
                })

        msg += "=" * 60 + "\n"
        if total_surplus > 0:
            msg += f"📊 Всего излишков: {total_surplus} шт.\n"
        if total_shortage > 0:
            msg += f"📊 Всего недостачи: {total_shortage} шт.\n"
        msg += "\nПосле корректировки остатки будут приведены к учётным значениям.\n"
        msg += "Продолжить утверждение акта?"

        # Спрашиваем подтверждение
        if not messagebox.askyesno("Корректировка остатков", msg):
            messagebox.showinfo("Отмена", "Акт не утверждён. Корректировка отменена.")
            return

        # Второе подтверждение
        if not messagebox.askyesno("Подтверждение",
                                   f"Утвердить акт о расхождениях {act_number}? Отменить это действие будет невозможно."):
            messagebox.showinfo("Отмена", "Акт не утверждён.")
            return

        # Выполняем корректировку
        conn = get_db_connection()
        cur = conn.cursor()

        for corr in corrections:
            item_type = corr['item_type']
            item_name = corr['item_name']
            surplus_qty = corr['surplus']
            shortage_qty = corr['shortage']

            if item_type == 'material':
                if surplus_qty > 0:
                    # Излишки: уменьшаем фактическое количество (списываем лишнее)
                    cur.execute("""
                        UPDATE materials 
                        SET actual_quantity = actual_quantity - %s
                        WHERE name = %s
                    """, (surplus_qty, item_name))
                    # Записываем движение
                    cur.execute("""
                        INSERT INTO movements (material_id, operation_type, quantity, reason, created_by)
                        SELECT id, 'Списание излишков', %s, 'Акт о расхождениях', %s
                        FROM materials WHERE name = %s
                    """, (surplus_qty, self.username, item_name))

                if shortage_qty > 0:
                    # Недостача: увеличиваем фактическое количество (добавляем недостающее)
                    cur.execute("""
                        UPDATE materials 
                        SET actual_quantity = actual_quantity + %s
                        WHERE name = %s
                    """, (shortage_qty, item_name))
                    # Записываем движение
                    cur.execute("""
                        INSERT INTO movements (material_id, operation_type, quantity, reason, created_by)
                        SELECT id, 'Оприходование недостачи', %s, 'Акт о расхождениях', %s
                        FROM materials WHERE name = %s
                    """, (shortage_qty, self.username, item_name))

            elif item_type == 'inventory':
                if surplus_qty > 0:
                    cur.execute("""
                        UPDATE inventory_items 
                        SET actual_quantity = actual_quantity - %s
                        WHERE name = %s AND status = 'in_use'
                    """, (surplus_qty, item_name))
                    cur.execute("""
                        INSERT INTO movements (inventory_item_id, operation_type, quantity, reason, created_by)
                        SELECT id, 'Списание излишков', %s, 'Акт о расхождениях', %s
                        FROM inventory_items WHERE name = %s AND status = 'in_use'
                    """, (surplus_qty, self.username, item_name))

                if shortage_qty > 0:
                    cur.execute("""
                        UPDATE inventory_items 
                        SET actual_quantity = actual_quantity + %s
                        WHERE name = %s AND status = 'in_use'
                    """, (shortage_qty, item_name))
                    cur.execute("""
                        INSERT INTO movements (inventory_item_id, operation_type, quantity, reason, created_by)
                        SELECT id, 'Оприходование недостачи', %s, 'Акт о расхождениях', %s
                        FROM inventory_items WHERE name = %s AND status = 'in_use'
                    """, (shortage_qty, self.username, item_name))

        # Обновляем статус акта
        cur.execute("UPDATE discrepancy_acts SET status = 'approved' WHERE id = %s", (act_id,))
        conn.commit()
        cur.close()
        conn.close()

        # Обновляем все списки
        self.load_discrepancy_acts_list()

        if hasattr(self, 'material_tree') and self.material_tree:
            try:
                self.load_materials_data()
                self.load_materials_accounting_data()
            except:
                pass
        if hasattr(self, 'inventory_tree') and self.inventory_tree:
            try:
                self.load_inventory_data()
                self.load_inventory_accounting_data()
            except:
                pass

        messagebox.showinfo("Успех", f"Акт о расхождениях {act_number} утверждён.\nОстатки успешно скорректированы.")

    def view_discrepancy_act_details(self):
        selected = self.discrepancy_acts_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите акт")
            return

        act_id = self.discrepancy_acts_tree.item(selected[0])['values'][0]

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, act_number, act_date, commission_id, total_surplus_quantity, total_surplus_sum,
                   total_shortage_quantity, total_shortage_sum, reason, status, created_by, created_at
            FROM discrepancy_acts WHERE id = %s
        """, (act_id,))
        act = cur.fetchone()

        # Исправленный запрос для получения деталей
        cur.execute("""
            SELECT 
                dai.item_type,
                COALESCE(nv.item_name, '') as item_name,
                CASE 
                    WHEN dai.item_type = 'inventory' THEN COALESCE(i.inventory_number, '')
                    ELSE ''
                END as inventory_number,
                CASE 
                    WHEN dai.item_type = 'material' THEN COALESCE(m.location, '')
                    WHEN dai.item_type = 'inventory' THEN COALESCE(i.location, '')
                    ELSE ''
                END as location,
                CASE 
                    WHEN dai.item_type = 'inventory' THEN COALESCE(rp.full_name, '')
                    ELSE ''
                END as responsible_person,
                dai.surplus_quantity,
                dai.shortage_quantity
            FROM discrepancy_act_items dai
            LEFT JOIN notifications_view nv ON nv.item_type = dai.item_type
            LEFT JOIN inventory_items i ON i.name = nv.item_name AND dai.item_type = 'inventory'
            LEFT JOIN materials m ON m.name = nv.item_name AND dai.item_type = 'material'
            LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
            WHERE dai.discrepancy_act_id = %s
        """, (act_id,))
        items = cur.fetchall()
        cur.close()
        conn.close()

        if not act:
            messagebox.showerror("Ошибка", "Акт не найден")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Акт о расхождениях - {act[1]}")
        dialog.geometry("1000x700")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.focus_set()

        main_frame = tk.Frame(dialog, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        tk.Label(main_frame, text="АКТ О РАСХОЖДЕНИЯХ", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10)

        info_left = tk.Frame(info_frame, bg='white')
        info_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=5)
        info_right = tk.Frame(info_frame, bg='white')
        info_right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=5)

        tk.Label(info_left, text=f"Номер акта: {act[1]}", font=('Arial', 10), bg='white', anchor='w').pack(anchor='w',
                                                                                                           pady=2)
        tk.Label(info_left, text=f"Дата: {act[2]}", font=('Arial', 10), bg='white', anchor='w').pack(anchor='w', pady=2)
        tk.Label(info_left, text=f"Статус: {act[9]}", font=('Arial', 10, 'bold'), bg='white',
                 fg='#27ae60' if act[9] == 'approved' else '#e74c3c', anchor='w').pack(anchor='w', pady=2)
        tk.Label(info_right, text=f"Создал: {act[10]}", font=('Arial', 10), bg='white', anchor='w').pack(anchor='w',
                                                                                                         pady=2)
        tk.Label(info_right, text=f"Причина: {act[8] if act[8] else 'не указана'}", font=('Arial', 10), bg='white',
                 anchor='w', wraplength=400).pack(anchor='w', pady=2)

        items_frame = tk.LabelFrame(main_frame, text="Выявленные расхождения по каждому объекту",
                                    font=('Arial', 12, 'bold'), bg='white')
        items_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        columns = ('item_type', 'item_name', 'inventory_number', 'location', 'responsible_person',
                   'surplus_quantity', 'shortage_quantity')
        tree = ttk.Treeview(items_frame, columns=columns, show='headings', height=10)

        headings = {'item_type': 'Тип', 'item_name': 'Наименование', 'inventory_number': 'Инв.номер',
                    'location': 'Место', 'responsible_person': 'МОЛ',
                    'surplus_quantity': 'Излишки (кол-во)', 'shortage_quantity': 'Недостача (кол-во)'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=120 if col not in ['item_name', 'responsible_person'] else 150)

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        total_surplus = 0
        total_shortage = 0
        for item in items:
            item_type, item_name, inv_num, location, responsible_person, surplus_qty, shortage_qty = item
            tree.insert('', tk.END,
                        values=(item_type, item_name, inv_num, location, responsible_person, surplus_qty, shortage_qty))
            total_surplus += surplus_qty if surplus_qty else 0
            total_shortage += shortage_qty if shortage_qty else 0

        total_frame = tk.Frame(main_frame, bg='white')
        total_frame.pack(fill=tk.X, pady=10)

        total_surplus_sum = act[5] if act[5] else 0
        total_shortage_sum = act[7] if act[7] else 0

        tk.Label(total_frame, text=f"Всего излишков: {total_surplus} шт. ({total_surplus_sum:,.2f} руб.)",
                 font=('Arial', 10, 'bold'), bg='white', fg='#e67e22').pack()
        tk.Label(total_frame, text=f"Всего недостачи: {total_shortage} шт. ({total_shortage_sum:,.2f} руб.)",
                 font=('Arial', 10, 'bold'), bg='white', fg='#e74c3c').pack()

        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=15)

        if act[9] != 'approved' and self.role in ['admin', 'storekeeper']:
            def approve_this_act():
                if messagebox.askyesno("Подтверждение", "Утвердить акт?"):
                    conn2 = get_db_connection()
                    cur2 = conn2.cursor()
                    cur2.execute("UPDATE discrepancy_acts SET status = 'approved' WHERE id = %s", (act_id,))
                    conn2.commit()
                    cur2.close()
                    conn2.close()
                    messagebox.showinfo("Успех", "Акт утверждён")
                    dialog.destroy()
                    self.show_discrepancy_acts()

            tk.Button(btn_frame, text="✅ Утвердить акт", command=approve_this_act, bg='#1abc9c', fg='white',
                      width=15).pack(side=tk.LEFT, padx=10)

        def export_this_act():
            self.export_discrepancy_act(act_id)

        tk.Button(btn_frame, text="📎 Экспорт в Excel", command=export_this_act, bg='#9b59b6', fg='white',
                  width=15).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="✖️ Закрыть", command=dialog.destroy, bg='#95a5a6', fg='white', width=15).pack(
            side=tk.LEFT, padx=10)

    def export_discrepancy_act(self, act_id=None):
        if act_id is None:
            selected = self.discrepancy_acts_tree.selection()
            if not selected:
                messagebox.showwarning("Ошибка", "Выберите акт")
                return
            act_id = self.discrepancy_acts_tree.item(selected[0])['values'][0]

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT act_number, act_date, total_surplus_quantity, total_surplus_sum,
                   total_shortage_quantity, total_shortage_sum, reason, created_by
            FROM discrepancy_acts WHERE id = %s
        """, (act_id,))
        act = cur.fetchone()

        cur.execute("""
            SELECT dai.item_type, dai.surplus_quantity, dai.shortage_quantity,
                   COALESCE(nv.item_name, '') as item_name,
                   CASE 
                       WHEN dai.item_type = 'inventory' THEN COALESCE(i.inventory_number, '')
                       ELSE ''
                   END as inventory_number
            FROM discrepancy_act_items dai
            LEFT JOIN notifications_view nv ON nv.item_type = dai.item_type
            LEFT JOIN inventory_items i ON i.name = nv.item_name AND dai.item_type = 'inventory'
            WHERE dai.discrepancy_act_id = %s
        """, (act_id,))
        items = cur.fetchall()
        cur.close()
        conn.close()

        if not act:
            messagebox.showerror("Ошибка", "Акт не найден")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel документы", "*.xlsx")],
            initialfile=f"Акт_расхождений_{act[0]}.xlsx"
        )
        if not file_path:
            return

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Акт о расхождениях"

            ws['A1'] = "АКТ О РАСХОЖДЕНИЯХ"
            ws.merge_cells('A1:F1')
            ws['A1'].font = Font(size=16, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')

            ws['A3'] = f"Номер акта: {act[0]}"
            ws['A4'] = f"Дата: {act[1]}"
            ws['A5'] = f"Создал: {act[7]}"
            ws['A6'] = f"Причина: {act[6] if act[6] else 'не указана'}"

            headers = ['Тип', 'Наименование', 'Инв.номер', 'Излишки (кол-во)', 'Недостача (кол-во)']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=8, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')

            for i, item in enumerate(items, 9):
                ws.cell(row=i, column=1, value=item[0])
                ws.cell(row=i, column=2, value=item[3])
                ws.cell(row=i, column=3, value=item[4])
                ws.cell(row=i, column=4, value=item[1])
                ws.cell(row=i, column=5, value=item[2])

            last_row = 9 + len(items)
            ws.cell(row=last_row + 1, column=1, value=f"ИТОГО излишков: {act[2]} шт. ({act[3]:,.2f} руб.)")
            ws.cell(row=last_row + 2, column=1, value=f"ИТОГО недостачи: {act[4]} шт. ({act[5]:,.2f} руб.)")

            for col in range(1, 6):
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 25

            wb.save(file_path)
            messagebox.showinfo("Успех", f"Акт сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def create_discrepancy_act(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT item_type, item_name, surplus_quantity, surplus_sum, shortage_quantity, shortage_sum
            FROM notifications_view
            WHERE surplus_quantity != 0 OR shortage_quantity != 0
        """)
        discrepancies = cur.fetchall()
        cur.close()
        conn.close()

        if not discrepancies:
            messagebox.showinfo("Информация", "Нет расхождений для создания акта")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Создание акта о расхождениях")
        dialog.geometry("900x700")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.focus_set()

        main_frame = tk.Frame(dialog, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        tk.Label(main_frame, text="АКТ О РАСХОЖДЕНИЯХ", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COALESCE(MAX(id), 0) + 1 FROM discrepancy_acts")
        next_id = cur.fetchone()[0]
        cur.close()
        conn.close()
        act_number = f"АР-{next_id:04d}"

        tk.Label(info_frame, text="Номер акта:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5,
                                                                                      sticky='w')
        tk.Label(info_frame, text=act_number, font=('Arial', 10, 'bold'), bg='white', fg='#27ae60').grid(row=0,
                                                                                                         column=1,
                                                                                                         padx=10,
                                                                                                         pady=5,
                                                                                                         sticky='w')

        tk.Label(info_frame, text="Дата акта:", font=('Arial', 10), bg='white').grid(row=0, column=2, padx=10, pady=5,
                                                                                     sticky='w')
        date_entry = tk.Entry(info_frame, width=15)
        date_entry.insert(0, datetime.now().strftime('%Y-%m-%d'))
        date_entry.grid(row=0, column=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Комиссия:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10, pady=5,
                                                                                    sticky='w')
        commission = self.get_commission()
        commission_label = tk.Label(info_frame, text=commission[1] if commission else "Не назначена",
                                    font=('Arial', 10), bg='white')
        commission_label.grid(row=1, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        items_frame = tk.LabelFrame(main_frame, text="Выявленные расхождения", font=('Arial', 12, 'bold'), bg='white')
        items_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        # Добавляем колонку "Объект" для отображения деталей
        columns = ('type', 'name', 'object_info', 'surplus_qty', 'surplus_sum', 'shortage_qty', 'shortage_sum')
        items_tree = ttk.Treeview(items_frame, columns=columns, show='headings', height=10)

        headings = {'type': 'Тип', 'name': 'Наименование', 'object_info': 'Объект (инв.номер/место/МОЛ)',
                    'surplus_qty': 'Излишки (кол-во)', 'surplus_sum': 'Излишки (сумма)',
                    'shortage_qty': 'Недостача (кол-во)', 'shortage_sum': 'Недостача (сумма)'}
        for col, heading in headings.items():
            items_tree.heading(col, text=heading)
            items_tree.column(col, width=120 if col != 'object_info' else 250)

        items_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Сохраняем полные данные о расхождениях для последующего сохранения
        discrepancies_full = []
        for disc in discrepancies:
            item_type, item_name, surplus_qty, surplus_sum, shortage_qty, shortage_sum = disc

            # Получаем дополнительную информацию об объекте
            object_info = ""
            inv_number = ""
            location = ""
            responsible_person = ""

            if item_type == 'inventory':
                conn2 = get_db_connection()
                cur2 = conn2.cursor()
                cur2.execute("""
                    SELECT inventory_number, location, rp.full_name
                    FROM inventory_items i
                    LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
                    WHERE i.name = %s AND i.status = 'in_use'
                """, (item_name,))
                inv_data = cur2.fetchone()
                cur2.close()
                conn2.close()
                if inv_data:
                    inv_number = inv_data[0] if inv_data[0] else ''
                    location = inv_data[1] if inv_data[1] else ''
                    responsible_person = inv_data[2] if inv_data[2] else ''
                    object_info = f"инв.№{inv_number}, место: {location}, МОЛ: {responsible_person}"
            else:
                conn2 = get_db_connection()
                cur2 = conn2.cursor()
                cur2.execute("SELECT location FROM materials WHERE name = %s", (item_name,))
                mat_data = cur2.fetchone()
                cur2.close()
                conn2.close()
                if mat_data:
                    location = mat_data[0] if mat_data[0] else ''
                    object_info = f"место: {location}"

            items_tree.insert('', tk.END, values=(
                item_type, item_name, object_info, surplus_qty, surplus_sum, shortage_qty, shortage_sum
            ))
            discrepancies_full.append({
                'item_type': item_type,
                'item_name': item_name,
                'surplus_quantity': surplus_qty,
                'shortage_quantity': shortage_qty,
                'inventory_number': inv_number,
                'location': location,
                'responsible_person': responsible_person
            })

        total_surplus_qty = sum(d[2] for d in discrepancies)
        total_surplus_sum = sum(d[3] for d in discrepancies)
        total_shortage_qty = sum(d[4] for d in discrepancies)
        total_shortage_sum = sum(d[5] for d in discrepancies)

        total_frame = tk.Frame(main_frame, bg='white')
        total_frame.pack(fill=tk.X, pady=10)
        total_text = f"Всего излишков: {total_surplus_qty} шт. ({total_surplus_sum:,.2f} руб.) | Всего недостачи: {total_shortage_qty} шт. ({total_shortage_sum:,.2f} руб.)"
        tk.Label(total_frame, text=total_text, font=('Arial', 10, 'bold'), bg='white', fg='#e74c3c').pack()

        tk.Label(info_frame, text="Причина расхождений:", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10,
                                                                                               pady=5, sticky='w')
        reason_entry = tk.Entry(info_frame, width=40)
        reason_entry.grid(row=2, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        def save_act():
            try:
                act_date = date_entry.get().strip()
                reason = reason_entry.get().strip()

                conn2 = get_db_connection()
                cur2 = conn2.cursor()

                commission_id = commission[0] if commission else None

                cur2.execute("""
                    INSERT INTO discrepancy_acts (act_number, act_date, commission_id, 
                          total_surplus_quantity, total_surplus_sum, total_shortage_quantity, total_shortage_sum,
                          reason, status, created_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'draft', %s)
                    RETURNING id
                """, (act_number, act_date, commission_id, total_surplus_qty, total_surplus_sum,
                      total_shortage_qty, total_shortage_sum, reason, self.username))
                act_id = cur2.fetchone()[0]

                for disc in discrepancies_full:
                    cur2.execute("""
                        INSERT INTO discrepancy_act_items (discrepancy_act_id, item_type,
                              surplus_quantity, shortage_quantity)
                        VALUES (%s, %s, %s, %s)
                    """, (act_id, disc['item_type'], disc['surplus_quantity'], disc['shortage_quantity']))

                conn2.commit()
                cur2.close()
                conn2.close()

                messagebox.showinfo("Успех", f"Акт {act_number} создан в статусе 'черновик'")
                dialog.destroy()
                self.show_discrepancy_acts()
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=15)
        tk.Button(btn_frame, text="💾 Сохранить как черновик", command=save_act, bg='#f39c12', fg='white',
                  width=20).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="✖️ Отмена", command=dialog.destroy, bg='#95a5a6', fg='white', width=15).pack(
            side=tk.LEFT, padx=10)

    def show_reports(self):
        self.highlight_menu_button("📈 Отчёты")
        self.clear_content()

        tk.Label(self.content_frame, text="Отчёты", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        btn_frame.pack(pady=10)

        tk.Button(btn_frame, text="📊 Текущие остатки (материалы)", command=self.report_materials_stock,
                  bg='#3498db', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="🪑 Текущие остатки (ОС)", command=self.report_inventory_stock,
                  bg='#3498db', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="📋 История движений", command=self.report_movements,
                  bg='#3498db', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="⚠️ Мало на складе", command=self.report_low_stock,
                  bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="🔄 Расхождения", command=self.report_discrepancies,
                  bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="📄 Инвентарная ведомость по МОЛ", command=self.show_inventory_by_person,
                  bg='#9b59b6', fg='white').pack(side=tk.LEFT, padx=10)

    def report_materials_stock(self):
        self.clear_content()

        tk.Label(self.content_frame, text="Текущие остатки материалов", font=('Arial', 16, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        filter_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Фильтр по месту:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        location_filter = ttk.Combobox(filter_frame, values=[''] + self.get_locations('material'), width=25)
        location_filter.pack(side=tk.LEFT, padx=5)

        tk.Label(filter_frame, text="Фильтр по категории:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        category_filter = ttk.Combobox(filter_frame, values=[''] + self.get_categories(), width=25)
        category_filter.pack(side=tk.LEFT, padx=5)

        def apply_filters():
            loc = location_filter.get()
            cat = category_filter.get()
            for item in tree.get_children():
                tree.delete(item)

            conn = get_db_connection()
            cur = conn.cursor()
            query = """
                SELECT m.name, COALESCE(c.name, '') as category, m.actual_quantity, COALESCE(u.short_name, '') as unit, 
                       m.price, COALESCE(m.location, '') as location, m.actual_sum
                FROM materials m
                LEFT JOIN asset_categories c ON m.category_id = c.id
                LEFT JOIN units u ON m.unit_id = u.id
                WHERE m.actual_quantity > 0
            """
            params = []
            if loc:
                query += " AND m.location = %s"
                params.append(loc)
            if cat:
                query += " AND c.name = %s"
                params.append(cat)
            query += " ORDER BY m.name"
            cur.execute(query, params)
            for row in cur.fetchall():
                tree.insert('', tk.END, values=row)
            cur.close()
            conn.close()

        tk.Button(filter_frame, text="Применить фильтр", command=apply_filters, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(filter_frame, text="Экспорт в Excel", command=lambda: self.export_report_to_excel(tree, "Остатки_материалов"), bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=10)

        columns = ('name', 'category', 'actual_quantity', 'unit', 'price', 'location', 'actual_sum')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'name': 'Наименование', 'category': 'Категория', 'actual_quantity': 'Кол-во',
                    'unit': 'Ед.изм', 'price': 'Цена', 'location': 'Место', 'actual_sum': 'Сумма'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150 if col != 'name' else 250)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        apply_filters()

    def report_inventory_stock(self):
        self.clear_content()

        tk.Label(self.content_frame, text="Основные средства на балансе", font=('Arial', 16, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        filter_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Фильтр по МОЛ:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        person_filter = ttk.Combobox(filter_frame, values=[''] + self.get_persons(), width=25)
        person_filter.pack(side=tk.LEFT, padx=5)

        tk.Label(filter_frame, text="Фильтр по месту:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        location_filter = ttk.Combobox(filter_frame, values=[''] + self.get_locations('inventory'), width=25)
        location_filter.pack(side=tk.LEFT, padx=5)

        def apply_filters():
            pers = person_filter.get()
            loc = location_filter.get()
            for item in tree.get_children():
                tree.delete(item)

            conn = get_db_connection()
            cur = conn.cursor()
            query = """
                SELECT i.inventory_number, i.name, COALESCE(c.name, '') as category, i.actual_quantity, 
                       i.price, COALESCE(i.location, '') as location, COALESCE(rp.full_name, '') as responsible_person,
                       i.actual_sum
                FROM inventory_items i
                LEFT JOIN asset_categories c ON i.category_id = c.id
                LEFT JOIN responsible_persons rp ON i.responsible_person_id = rp.id
                WHERE i.status = 'in_use' AND i.actual_quantity > 0
            """
            params = []
            if pers:
                query += " AND rp.full_name = %s"
                params.append(pers)
            if loc:
                query += " AND i.location = %s"
                params.append(loc)
            query += " ORDER BY i.name"
            cur.execute(query, params)
            for row in cur.fetchall():
                tree.insert('', tk.END, values=row)
            cur.close()
            conn.close()

        tk.Button(filter_frame, text="Применить фильтр", command=apply_filters, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(filter_frame, text="Экспорт в Excel", command=lambda: self.export_report_to_excel(tree, "Остатки_ОС"), bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=10)

        columns = ('inventory_number', 'name', 'category', 'actual_quantity', 'price', 'location', 'responsible_person', 'actual_sum')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'inventory_number': 'Инв.номер', 'name': 'Наименование', 'category': 'Категория',
                    'actual_quantity': 'Кол-во', 'price': 'Цена', 'location': 'Место',
                    'responsible_person': 'МОЛ', 'actual_sum': 'Сумма'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150 if col != 'name' else 200)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        apply_filters()

    def report_movements(self):
        self.clear_content()

        tk.Label(self.content_frame, text="История движений", font=('Arial', 16, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        filter_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="С даты:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        from_date = tk.Entry(filter_frame, width=12)
        from_date.insert(0, (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
        from_date.pack(side=tk.LEFT, padx=5)

        tk.Label(filter_frame, text="По дату:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        to_date = tk.Entry(filter_frame, width=12)
        to_date.insert(0, datetime.now().strftime('%Y-%m-%d'))
        to_date.pack(side=tk.LEFT, padx=5)

        tk.Label(filter_frame, text="Тип операции:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        op_filter = ttk.Combobox(filter_frame, values=['', 'Приход', 'Выдача', 'Списание', 'Перемещение', 'Переход'], width=15)
        op_filter.pack(side=tk.LEFT, padx=5)

        def apply_filters():
            from_d = from_date.get()
            to_d = to_date.get()
            op_type = op_filter.get()
            for item in tree.get_children():
                tree.delete(item)

            op_map = {'Приход': 'receipt', 'Выдача': 'issue', 'Списание': 'write_off', 'Перемещение': 'move', 'Переход': 'Переход'}

            conn = get_db_connection()
            cur = conn.cursor()
            query = """
                SELECT 'Материал' as item_type, m.name, mov.operation_type, mov.quantity, rp.full_name, 
                       mov.operation_date, mov.created_by
                FROM movements mov
                LEFT JOIN materials m ON mov.material_id = m.id
                LEFT JOIN responsible_persons rp ON mov.person_id = rp.id
                WHERE mov.material_id IS NOT NULL AND mov.operation_date BETWEEN %s AND %s
                UNION ALL
                SELECT 'ОС' as item_type, i.name, mov.operation_type, mov.quantity, rp.full_name, 
                       mov.operation_date, mov.created_by
                FROM movements mov
                LEFT JOIN inventory_items i ON mov.inventory_item_id = i.id
                LEFT JOIN responsible_persons rp ON mov.person_id = rp.id
                WHERE mov.inventory_item_id IS NOT NULL AND mov.operation_date BETWEEN %s AND %s
            """
            params = [from_d, to_d, from_d, to_d]

            if op_type and op_type in op_map:
                query += " AND mov.operation_type = %s"
                params.append(op_map[op_type])

            query += " ORDER BY operation_date DESC LIMIT 500"

            cur.execute(query, params)
            for row in cur.fetchall():
                op_display = row[2]
                if op_display == 'receipt':
                    op_display = '📥 Приход'
                elif op_display == 'issue':
                    op_display = '📤 Выдача'
                elif op_display == 'write_off':
                    op_display = '❌ Списание'
                elif op_display == 'move':
                    op_display = '📦 Перемещение'
                elif op_display == 'Переход':
                    op_display = '📋 Переход'
                values = list(row)
                values[2] = op_display
                tree.insert('', tk.END, values=values)
            cur.close()
            conn.close()

        tk.Button(filter_frame, text="Применить фильтр", command=apply_filters, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(filter_frame, text="Экспорт в Excel", command=lambda: self.export_report_to_excel(tree, "История_движений"), bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=10)

        columns = ('item_type', 'item_name', 'operation_type', 'quantity', 'person', 'operation_date', 'created_by')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'item_type': 'Тип', 'item_name': 'Наименование', 'operation_type': 'Операция',
                    'quantity': 'Кол-во', 'person': 'Сотрудник', 'operation_date': 'Дата', 'created_by': 'Кто создал'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        apply_filters()

    def report_low_stock(self):
        self.clear_content()

        tk.Label(self.content_frame, text="Материалы с низким остатком", font=('Arial', 16, 'bold'),
                 bg='#ecf0f1', fg='#e74c3c').pack(pady=10)

        filter_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Фильтр по месту:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        location_filter = ttk.Combobox(filter_frame, values=[''] + self.get_locations('material'), width=25)
        location_filter.pack(side=tk.LEFT, padx=5)

        def apply_filters():
            loc = location_filter.get()
            for item in tree.get_children():
                tree.delete(item)

            conn = get_db_connection()
            cur = conn.cursor()
            query = """
                SELECT m.name, COALESCE(c.name, '') as category, m.actual_quantity, m.min_quantity, 
                       COALESCE(u.short_name, '') as unit, COALESCE(m.location, '') as location,
                       (m.min_quantity - m.actual_quantity) as deficit
                FROM materials m
                LEFT JOIN asset_categories c ON m.category_id = c.id
                LEFT JOIN units u ON m.unit_id = u.id
                WHERE m.actual_quantity <= m.min_quantity AND m.min_quantity > 0
            """
            params = []
            if loc:
                query += " AND m.location = %s"
                params.append(loc)
            query += " ORDER BY (m.actual_quantity / NULLIF(m.min_quantity, 0))"
            cur.execute(query, params)
            for row in cur.fetchall():
                tree.insert('', tk.END, values=row)
            cur.close()
            conn.close()

        tk.Button(filter_frame, text="Применить фильтр", command=apply_filters, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(filter_frame, text="Экспорт в Excel", command=lambda: self.export_report_to_excel(tree, "Мало_на_складе"), bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=10)

        columns = ('name', 'category', 'actual_quantity', 'min_quantity', 'unit', 'location', 'deficit')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'name': 'Наименование', 'category': 'Категория', 'actual_quantity': 'Кол-во',
                    'min_quantity': 'Мин. остаток', 'unit': 'Ед.изм', 'location': 'Место', 'deficit': 'Не хватает'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150 if col != 'name' else 250)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        apply_filters()

    def report_discrepancies(self):
        self.clear_content()

        tk.Label(self.content_frame, text="Расхождения (излишки/недостача)", font=('Arial', 16, 'bold'),
                 bg='#ecf0f1', fg='#e74c3c').pack(pady=10)

        filter_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Фильтр по типу:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        type_filter = ttk.Combobox(filter_frame, values=['', 'Материалы', 'Основные средства'], width=20)
        type_filter.pack(side=tk.LEFT, padx=5)

        def apply_filters():
            typ = type_filter.get()
            for item in tree.get_children():
                tree.delete(item)

            conn = get_db_connection()
            cur = conn.cursor()
            query = """
                SELECT item_type, item_name, surplus_quantity, surplus_sum, shortage_quantity, shortage_sum
                FROM notifications_view
                WHERE surplus_quantity != 0 OR shortage_quantity != 0
            """
            params = []
            if typ == 'Материалы':
                query += " AND item_type = 'material'"
            elif typ == 'Основные средства':
                query += " AND item_type = 'inventory'"
            cur.execute(query, params)
            for row in cur.fetchall():
                tree.insert('', tk.END, values=row)
            cur.close()
            conn.close()

        tk.Button(filter_frame, text="Применить фильтр", command=apply_filters, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(filter_frame, text="Экспорт в Excel", command=lambda: self.export_report_to_excel(tree, "Расхождения"), bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=10)

        columns = ('item_type', 'item_name', 'surplus_quantity', 'surplus_sum', 'shortage_quantity', 'shortage_sum')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'item_type': 'Тип', 'item_name': 'Наименование',
                    'surplus_quantity': 'Излишки (кол-во)', 'surplus_sum': 'Излишки (сумма)',
                    'shortage_quantity': 'Недостача (кол-во)', 'shortage_sum': 'Недостача (сумма)'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        apply_filters()

    def export_report_to_excel(self, tree, filename_prefix):
        if not EXCEL_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека openpyxl не установлена")
            return

        columns = [tree.heading(col)['text'] for col in tree['columns']]
        data = []
        for child in tree.get_children():
            values = tree.item(child)['values']
            data.append(values)

        if not data:
            messagebox.showwarning("Нет данных", "Нет данных для экспорта")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel документы", "*.xlsx")],
            initialfile=f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        if not file_path:
            return

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = filename_prefix[:31]

            for col_idx, header in enumerate(columns, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')

            for row_idx, row in enumerate(data, 2):
                for col_idx, value in enumerate(row, 1):
                    ws.cell(row=row_idx, column=col_idx, value=value)

            for col_idx in range(1, len(columns) + 1):
                max_length = len(columns[col_idx - 1])
                for row_idx in range(2, len(data) + 2):
                    cell_value = ws.cell(row=row_idx, column=col_idx).value
                    if cell_value:
                        max_length = max(max_length, len(str(cell_value)))
                ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_length + 2, 40)

            wb.save(file_path)
            messagebox.showinfo("Успех", f"Отчёт сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def show_users(self):
        if self.role != 'admin':
            return
        self.highlight_menu_button("🔐 Пользователи")
        self.clear_content()

        tk.Label(self.content_frame, text="Управление пользователями", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        btn_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        btn_frame.pack(pady=5)

        tk.Button(btn_frame, text="➕ Добавить пользователя", command=self.add_user, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="✏️ Редактировать", command=self.edit_user, bg='#f39c12', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="🗑️ Удалить", command=self.delete_user, bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        search_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        search_frame.pack(pady=5, fill=tk.X)

        tk.Label(search_frame, text="🔍 Поиск:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        self.user_search = tk.Entry(search_frame, width=40)
        self.user_search.pack(side=tk.LEFT, padx=5)
        self.user_search.bind('<KeyRelease>', lambda e: self.load_users())
        tk.Button(search_frame, text="Очистить", command=self.clear_user_search, bg='#95a5a6', fg='white').pack(side=tk.LEFT, padx=5)

        columns = ('id', 'username', 'role', 'full_name')
        self.user_tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=15)

        headings = {'id': 'ID', 'username': 'Логин', 'role': 'Роль', 'full_name': 'ФИО'}
        for col, heading in headings.items():
            self.user_tree.heading(col, text=heading)
            self.user_tree.column(col, width=180)
            self.user_tree.heading(col, command=lambda c=col: self.sort_treeview(self.user_tree, c, False))

        self.user_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        self.load_users()

    def load_users(self):
        for item in self.user_tree.get_children():
            self.user_tree.delete(item)

        search_text = self.user_search.get().strip().lower()

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT u.id, u.username, u.role, COALESCE(rp.full_name, '') as full_name
            FROM users u
            LEFT JOIN responsible_persons rp ON u.person_id = rp.id
            ORDER BY u.id
        """)
        for row in cur.fetchall():
            if search_text == "" or search_text in str(row[1]).lower():
                self.user_tree.insert('', tk.END, values=row)
        cur.close()
        conn.close()

    def clear_user_search(self):
        self.user_search.delete(0, tk.END)
        self.load_users()

    def add_user(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Добавить пользователя")
        dialog.geometry("450x500")

        tk.Label(dialog, text="Логин:").grid(row=0, column=0, padx=10, pady=5, sticky='w')
        username_entry = tk.Entry(dialog, width=30)
        username_entry.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Пароль:").grid(row=1, column=0, padx=10, pady=5, sticky='w')
        password_entry = tk.Entry(dialog, width=30, show='*')
        password_entry.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Подтверждение пароля:").grid(row=2, column=0, padx=10, pady=5, sticky='w')
        confirm_entry = tk.Entry(dialog, width=30, show='*')
        confirm_entry.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Роль:").grid(row=3, column=0, padx=10, pady=5, sticky='w')
        role_combo = ttk.Combobox(dialog, values=['admin', 'storekeeper', 'accountant'], width=27)
        role_combo.grid(row=3, column=1, padx=10, pady=5)
        role_combo.current(1)

        tk.Label(dialog, text="Связать с сотрудником:").grid(row=4, column=0, padx=10, pady=5, sticky='w')
        persons = self.get_persons()
        person_combo = ttk.Combobox(dialog, values=[''] + persons, width=27)
        person_combo.grid(row=4, column=1, padx=10, pady=5)

        tk.Label(dialog, text="ФИО (если не выбран сотрудник):").grid(row=5, column=0, padx=10, pady=5, sticky='w')
        full_name_entry = tk.Entry(dialog, width=30)
        full_name_entry.grid(row=5, column=1, padx=10, pady=5)

        def save():
            username = username_entry.get().strip()
            password = password_entry.get().strip()
            confirm = confirm_entry.get().strip()
            role = role_combo.get()
            person_name = person_combo.get()
            full_name = full_name_entry.get().strip()

            if not username:
                messagebox.showerror("Ошибка", "Введите логин")
                return
            if not password:
                messagebox.showerror("Ошибка", "Введите пароль")
                return
            if password != confirm:
                messagebox.showerror("Ошибка", "Пароли не совпадают")
                return
            if not person_name and not full_name:
                messagebox.showerror("Ошибка", "Выберите сотрудника или введите ФИО")
                return

            conn = get_db_connection()
            cur = conn.cursor()

            try:
                person_id = None
                if person_name:
                    cur.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (person_name,))
                    person = cur.fetchone()
                    if person:
                        person_id = person[0]
                    else:
                        cur.execute("INSERT INTO responsible_persons (full_name) VALUES (%s) RETURNING id", (person_name,))
                        person_id = cur.fetchone()[0]
                elif full_name:
                    cur.execute("INSERT INTO responsible_persons (full_name) VALUES (%s) RETURNING id", (full_name,))
                    person_id = cur.fetchone()[0]

                password_hash = hash_password(password)
                cur.execute("""
                    INSERT INTO users (username, password_hash, role, person_id, is_active)
                    VALUES (%s, %s, %s, %s, TRUE)
                """, (username, password_hash, role, person_id))

                conn.commit()
                self.log_action(f"Добавлен пользователь {username} с ролью {role}")
                dialog.destroy()
                self.load_users()
                messagebox.showinfo("Успех", f"Пользователь {username} добавлен")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))
            finally:
                cur.close()
                conn.close()

        tk.Button(dialog, text="Сохранить", command=save, bg='#27ae60', fg='white').grid(row=6, column=0, columnspan=2, pady=20)

    def edit_user(self):
        selected = self.user_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите пользователя")
            return

        values = self.user_tree.item(selected[0])['values']
        user_id = values[0]
        old_username = values[1]
        old_role = values[2]
        old_full_name = values[3] if values[3] else ''

        dialog = tk.Toplevel(self.root)
        dialog.title("Редактировать пользователя")
        dialog.geometry("450x450")

        tk.Label(dialog, text="Логин:").grid(row=0, column=0, padx=10, pady=5, sticky='w')
        username_entry = tk.Entry(dialog, width=30)
        username_entry.insert(0, old_username)
        username_entry.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Новый пароль (оставьте пустым, чтобы не менять):").grid(row=1, column=0, padx=10, pady=5, sticky='w')
        password_entry = tk.Entry(dialog, width=30, show='*')
        password_entry.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Роль:").grid(row=2, column=0, padx=10, pady=5, sticky='w')
        role_combo = ttk.Combobox(dialog, values=['admin', 'storekeeper', 'accountant'], width=27)
        role_combo.set(old_role)
        role_combo.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="ФИО:").grid(row=3, column=0, padx=10, pady=5, sticky='w')
        full_name_entry = tk.Entry(dialog, width=30)
        full_name_entry.insert(0, old_full_name)
        full_name_entry.grid(row=3, column=1, padx=10, pady=5)

        def save():
            username = username_entry.get().strip()
            role = role_combo.get()
            full_name = full_name_entry.get().strip()
            new_password = password_entry.get().strip()

            if not username:
                messagebox.showerror("Ошибка", "Введите логин")
                return

            conn = get_db_connection()
            cur = conn.cursor()

            try:
                if new_password:
                    password_hash = hash_password(new_password)
                    cur.execute("UPDATE users SET username=%s, password_hash=%s, role=%s WHERE id=%s",
                                (username, password_hash, role, user_id))
                else:
                    cur.execute("UPDATE users SET username=%s, role=%s WHERE id=%s",
                                (username, role, user_id))

                if full_name:
                    cur.execute("""
                        UPDATE responsible_persons SET full_name=%s 
                        WHERE id = (SELECT person_id FROM users WHERE id=%s)
                    """, (full_name, user_id))

                conn.commit()
                self.log_action(f"Изменён пользователь {username}")
                dialog.destroy()
                self.load_users()
                messagebox.showinfo("Успех", f"Пользователь {username} обновлён")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))
            finally:
                cur.close()
                conn.close()

        tk.Button(dialog, text="Сохранить", command=save, bg='#f39c12', fg='white').grid(row=4, column=0, columnspan=2, pady=20)

    def delete_user(self):
        selected = self.user_tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите пользователя")
            return

        values = self.user_tree.item(selected[0])['values']
        user_id = values[0]
        username = values[1]

        if username == self.username:
            messagebox.showerror("Ошибка", "Нельзя удалить самого себя")
            return

        if messagebox.askyesno("Подтверждение", f"Удалить пользователя '{username}'?"):
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute("DELETE FROM users WHERE id = %s", (user_id,))
                conn.commit()
                self.log_action(f"Удалён пользователь {username}")
                self.load_users()
                messagebox.showinfo("Успех", f"Пользователь {username} удалён")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))
            finally:
                cur.close()
                conn.close()

    def show_logs(self):
        if self.role != 'admin':
            messagebox.showerror("Доступ запрещён", "Только администратор может просматривать логи")
            return
        self.highlight_menu_button("📜 Логи")
        self.clear_content()

        tk.Label(self.content_frame, text="Журнал действий", font=('Arial', 18, 'bold'),
                 bg='#ecf0f1').pack(pady=10)

        filter_frame = tk.Frame(self.content_frame, bg='#ecf0f1')
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Label(filter_frame, text="Фильтр по пользователю:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        user_filter = ttk.Combobox(filter_frame, values=[''] + [row[0] for row in self.get_all_users()], width=20)
        user_filter.pack(side=tk.LEFT, padx=5)

        tk.Label(filter_frame, text="Фильтр по действию:", bg='#ecf0f1').pack(side=tk.LEFT, padx=5)
        action_filter = ttk.Combobox(filter_frame, values=['', 'LOGIN', 'LOGOUT', 'ADMIN_ACTION', 'CREATE', 'EDIT', 'DELETE'], width=15)
        action_filter.pack(side=tk.LEFT, padx=5)

        def apply_logs_filter():
            for item in tree.get_children():
                tree.delete(item)

            conn = get_db_connection()
            cur = conn.cursor()
            query = "SELECT username, action, object_type, details, created_at FROM action_logs WHERE 1=1"
            params = []
            if user_filter.get():
                query += " AND username = %s"
                params.append(user_filter.get())
            if action_filter.get():
                query += " AND action = %s"
                params.append(action_filter.get())
            query += " ORDER BY created_at DESC LIMIT 500"
            cur.execute(query, params)
            for row in cur.fetchall():
                tree.insert('', tk.END, values=row)
            cur.close()
            conn.close()

        tk.Button(filter_frame, text="Применить фильтр", command=apply_logs_filter, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=10)
        tk.Button(filter_frame, text="Экспорт в Excel", command=lambda: self.export_report_to_excel(tree, "Журнал_действий"), bg='#1abc9c', fg='white').pack(side=tk.LEFT, padx=10)

        columns = ('username', 'action', 'object_type', 'details', 'created_at')
        tree = ttk.Treeview(self.content_frame, columns=columns, show='headings', height=20)

        headings = {'username': 'Пользователь', 'action': 'Действие', 'object_type': 'Тип',
                    'details': 'Подробности', 'created_at': 'Время'}
        for col, heading in headings.items():
            tree.heading(col, text=heading)
            tree.column(col, width=150 if col != 'details' else 300)
            tree.heading(col, command=lambda c=col: self.sort_treeview(tree, c, False))

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        apply_logs_filter()

    def get_all_users(self):
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT DISTINCT username FROM action_logs ORDER BY username")
        users = cur.fetchall()
        cur.close()
        conn.close()
        return users

    def add_receipt_item_dialog(self, parent, tree, item_type, units):
        """Диалог добавления позиции в акт оприходования"""
        dialog = tk.Toplevel(parent)
        dialog.title(f"Добавить {item_type}")
        dialog.geometry("500x550")
        dialog.configure(bg='white')
        dialog.transient(parent)
        dialog.grab_set()

        tk.Label(dialog, text="Наименование:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        name_entry = tk.Entry(dialog, width=40)
        name_entry.grid(row=0, column=1, padx=10, pady=5)

        if item_type == 'inventory':
            tk.Label(dialog, text="Инвентарный номер:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10, pady=5, sticky='w')
            inv_number_entry = tk.Entry(dialog, width=40)
            inv_number_entry.grid(row=1, column=1, padx=10, pady=5)
            row_offset = 2
        else:
            inv_number_entry = None
            row_offset = 1

        tk.Label(dialog, text="Единица измерения:", font=('Arial', 10), bg='white').grid(row=row_offset, column=0, padx=10, pady=5, sticky='w')
        unit_combo = ttk.Combobox(dialog, values=units, width=37)
        unit_combo.grid(row=row_offset, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Количество (факт):", font=('Arial', 10), bg='white').grid(row=row_offset + 1, column=0, padx=10, pady=5, sticky='w')
        quantity_entry = tk.Entry(dialog, width=40)
        quantity_entry.grid(row=row_offset + 1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Количество (учёт):", font=('Arial', 10), bg='white').grid(row=row_offset + 2, column=0, padx=10, pady=5, sticky='w')
        accounting_quantity_entry = tk.Entry(dialog, width=40)
        accounting_quantity_entry.grid(row=row_offset + 2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Цена за единицу:", font=('Arial', 10), bg='white').grid(row=row_offset + 3, column=0, padx=10, pady=5, sticky='w')
        price_entry = tk.Entry(dialog, width=40)
        price_entry.grid(row=row_offset + 3, column=1, padx=10, pady=5)

        def add():
            name = name_entry.get().strip()
            unit = unit_combo.get().strip()
            quantity = quantity_entry.get().strip()
            accounting_quantity = accounting_quantity_entry.get().strip()
            price = price_entry.get().strip()

            if not name:
                messagebox.showerror("Ошибка", "Введите наименование")
                return
            if not quantity:
                messagebox.showerror("Ошибка", "Введите количество")
                return
            if not price:
                messagebox.showerror("Ошибка", "Введите цену")
                return

            try:
                qty = float(quantity)
                acc_qty = float(accounting_quantity) if accounting_quantity else qty
                prc = float(price)
                total = qty * prc

                if item_type == 'inventory' and inv_number_entry:
                    inv_num = inv_number_entry.get().strip()
                    if inv_num:
                        name = f"{name} (инв.№{inv_num})"

                tree.insert('', tk.END, values=(item_type, name, unit, qty, acc_qty, prc, total))
                self.update_tree_total(tree)
                dialog.destroy()
            except ValueError:
                messagebox.showerror("Ошибка", "Количество и цена должны быть числами")

        tk.Button(dialog, text="Добавить", command=add, bg='#27ae60', fg='white', width=15).grid(row=row_offset + 4, column=0, columnspan=2, pady=20)

    def add_write_off_item_dialog(self, parent, tree, item_type, units):
        """Диалог добавления позиции в акт списания"""
        dialog = tk.Toplevel(parent)
        dialog.title("Добавить позицию")
        dialog.geometry("500x450")
        dialog.configure(bg='white')
        dialog.transient(parent)
        dialog.grab_set()

        tk.Label(dialog, text="Наименование:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        name_entry = tk.Entry(dialog, width=40)
        name_entry.grid(row=0, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Единица измерения:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10, pady=5, sticky='w')
        unit_combo = ttk.Combobox(dialog, values=units, width=37)
        unit_combo.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Количество:", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10, pady=5, sticky='w')
        quantity_entry = tk.Entry(dialog, width=40)
        quantity_entry.grid(row=2, column=1, padx=10, pady=5)

        tk.Label(dialog, text="Цена за единицу:", font=('Arial', 10), bg='white').grid(row=3, column=0, padx=10, pady=5, sticky='w')
        price_entry = tk.Entry(dialog, width=40)
        price_entry.grid(row=3, column=1, padx=10, pady=5)

        def add():
            name = name_entry.get().strip()
            unit = unit_combo.get().strip()
            quantity = quantity_entry.get().strip()
            price = price_entry.get().strip()

            if not name:
                messagebox.showerror("Ошибка", "Введите наименование")
                return
            if not quantity:
                messagebox.showerror("Ошибка", "Введите количество")
                return
            if not price:
                messagebox.showerror("Ошибка", "Введите цену")
                return

            try:
                qty = float(quantity)
                prc = float(price)
                total = qty * prc

                tree.insert('', tk.END, values=(name, unit, qty, prc, total))
                self.update_tree_total(tree)
                dialog.destroy()
            except ValueError:
                messagebox.showerror("Ошибка", "Количество и цена должны быть числами")

        tk.Button(dialog, text="Добавить", command=add, bg='#27ae60', fg='white', width=15).grid(row=4, column=0, columnspan=2, pady=20)

    def delete_selected_tree_item(self, tree):
        """Удаление выбранной позиции из дерева и обновление итога"""
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите позицию")
            return

        tree.delete(selected[0])

        total = 0
        for item in tree.get_children():
            values = tree.item(item)['values']
            if len(values) >= 6 and values[5]:
                try:
                    total += float(values[5])
                except:
                    pass
            elif len(values) >= 5 and values[4]:
                try:
                    total += float(values[4])
                except:
                    pass

        def find_and_update_total_label(widget):
            for child in widget.winfo_children():
                if isinstance(child, tk.Label):
                    if 'ИТОГО:' in child.cget('text'):
                        child.config(text=f"{total:,.2f} руб.")
                        return True
                if find_and_update_total_label(child):
                    return True
            return False

        root_window = tree.winfo_toplevel()
        find_and_update_total_label(root_window)

    def update_tree_total(self, tree):
        """Обновление итоговой суммы в диалоговом окне"""
        total = 0
        for item in tree.get_children():
            values = tree.item(item)['values']
            if len(values) >= 6 and values[5]:
                try:
                    total += float(values[5])
                except:
                    pass
            elif len(values) >= 5 and values[4]:
                try:
                    total += float(values[4])
                except:
                    pass

        root_window = tree.winfo_toplevel()

        def find_label(widget):
            for child in widget.winfo_children():
                if isinstance(child, tk.Label):
                    if 'ИТОГО:' in child.cget('text'):
                        child.config(text=f"{total:,.2f} руб.")
                        return True
                if find_label(child):
                    return True
            return False

        find_label(root_window)

    def create_receipt_act(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Создание акта оприходования")
        dialog.geometry("1000x800")
        dialog.configure(bg='white')
        dialog.transient(self.root)  # Связываем с главным окном
        dialog.grab_set()  # Делаем модальным (блокируем другие окна)
        dialog.focus_set()  # Устанавливаем фокус
        #  ОСНОВНОЙ КОНТЕЙНЕР 
        main_container = tk.Frame(dialog, bg='white')
        main_container.pack(fill=tk.BOTH, expand=True)

        #  СКРОЛЛИНГ ДЛЯ ОСНОВНОГО СОДЕРЖИМОГО 
        canvas = tk.Canvas(main_container, bg='white')
        scrollbar = tk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = scrollable_frame

        tk.Label(main_frame, text="АКТ ОПРИХОДОВАНИЯ", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        #  ИНФОРМАЦИЯ ОБ АКТЕ 
        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10, padx=10)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COALESCE(MAX(id), 0) + 1 FROM receipt_acts")
        next_id = cur.fetchone()[0]
        cur.close()
        conn.close()
        act_number = f"АО-{next_id:04d}"

        tk.Label(info_frame, text="Номер акта:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        tk.Label(info_frame, text=act_number, font=('Arial', 10, 'bold'), bg='white', fg='#27ae60').grid(row=0, column=1, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Дата акта:", font=('Arial', 10), bg='white').grid(row=0, column=2, padx=10, pady=5, sticky='w')
        date_entry = tk.Entry(info_frame, width=15)
        date_entry.insert(0, datetime.now().strftime('%Y-%m-%d'))
        date_entry.grid(row=0, column=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Поставщик:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10, pady=5, sticky='w')
        suppliers = self.get_suppliers()
        supplier_combo = ttk.Combobox(info_frame, values=suppliers, width=40)
        supplier_combo.set('')
        supplier_combo.grid(row=1, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Тип поступления:", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10, pady=5, sticky='w')
        receipt_type_combo = ttk.Combobox(info_frame, values=['покупка', 'безвозмездно'], width=40)
        receipt_type_combo.current(0)
        receipt_type_combo.grid(row=2, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Ответственное лицо:", font=('Arial', 10), bg='white').grid(row=3, column=0, padx=10, pady=5, sticky='w')
        persons = self.get_persons()
        person_combo = ttk.Combobox(info_frame, values=persons, width=40)
        person_combo.grid(row=3, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        #  ПОЗИЦИИ АКТА 
        items_frame = tk.LabelFrame(main_frame, text="Позиции оприходования", font=('Arial', 12, 'bold'), bg='white')
        items_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)

        # Форма для добавления позиции
        add_frame = tk.Frame(items_frame, bg='white', relief=tk.GROOVE, bd=2)
        add_frame.pack(fill=tk.X, padx=10, pady=10)

        tk.Label(add_frame, text="Тип:", font=('Arial', 10, 'bold'), bg='white').grid(row=0, column=0, padx=5, pady=5, sticky='w')
        type_combo = ttk.Combobox(add_frame, values=['Материал', 'Основное средство'], width=15)
        type_combo.current(0)
        type_combo.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(add_frame, text="Выбрать существующий:", font=('Arial', 10, 'bold'), bg='white').grid(row=0, column=2, padx=5, pady=5, sticky='w')
        existing_combo = ttk.Combobox(add_frame, width=30)
        existing_combo.grid(row=0, column=3, padx=5, pady=5)

        tk.Button(add_frame, text="Обновить",
                  command=lambda: self.update_existing_items_list(type_combo.get(), existing_combo),
                  bg='#3498db', fg='white', width=8).grid(row=0, column=4, padx=2, pady=5)

        tk.Label(add_frame, text="Новое наименование:", font=('Arial', 10, 'bold'), bg='white').grid(row=1, column=0, padx=5, pady=5, sticky='w')
        new_name_entry = tk.Entry(add_frame, width=40)
        new_name_entry.grid(row=1, column=1, columnspan=3, padx=5, pady=5, sticky='w')

        tk.Label(add_frame, text="Единица измерения:", font=('Arial', 10, 'bold'), bg='white').grid(row=2, column=0, padx=5, pady=5, sticky='w')
        units = self.get_units()
        unit_combo = ttk.Combobox(add_frame, values=units, width=37)
        unit_combo.grid(row=2, column=1, columnspan=3, padx=5, pady=5, sticky='w')

        tk.Label(add_frame, text="Количество (факт):", font=('Arial', 10, 'bold'), bg='white').grid(row=3, column=0, padx=5, pady=5, sticky='w')
        quantity_entry = tk.Entry(add_frame, width=15)
        quantity_entry.grid(row=3, column=1, padx=5, pady=5)

        tk.Label(add_frame, text="Количество (учёт):", font=('Arial', 10, 'bold'), bg='white').grid(row=3, column=2, padx=5, pady=5, sticky='w')
        accounting_quantity_entry = tk.Entry(add_frame, width=15)
        accounting_quantity_entry.grid(row=3, column=3, padx=5, pady=5)

        tk.Label(add_frame, text="Цена за единицу:", font=('Arial', 10, 'bold'), bg='white').grid(row=4, column=0, padx=5, pady=5, sticky='w')
        price_entry = tk.Entry(add_frame, width=15)
        price_entry.grid(row=4, column=1, padx=5, pady=5)

        tk.Label(add_frame, text="Инвентарный номер (для ОС):", font=('Arial', 10, 'bold'), bg='white').grid(row=4, column=2, padx=5, pady=5, sticky='w')
        inv_number_entry = tk.Entry(add_frame, width=20)
        inv_number_entry.grid(row=4, column=3, padx=5, pady=5)

        # ===== НОВЫЕ ПОЛЯ =====
        tk.Label(add_frame, text="Место хранения/нахождения:", font=('Arial', 10, 'bold'), bg='white').grid(row=5, column=0, padx=5, pady=5, sticky='w')
        locations = self.get_locations('both')
        location_combo = ttk.Combobox(add_frame, values=locations, width=37)
        location_combo.grid(row=5, column=1, columnspan=3, padx=5, pady=5, sticky='w')

        tk.Label(add_frame, text="МОЛ (для ОС):", font=('Arial', 10, 'bold'), bg='white').grid(row=6, column=0, padx=5, pady=5, sticky='w')
        responsible_person_combo = ttk.Combobox(add_frame, values=persons, width=37)
        responsible_person_combo.grid(row=6, column=1, columnspan=3, padx=5, pady=5, sticky='w')

        #  ФУНКЦИЯ ОБНОВЛЕНИЯ СПИСКА 
        def on_existing_select(event):
            selected = existing_combo.get()
            if selected != "➕ Новое" and selected in self.existing_items_data:
                data = self.existing_items_data[selected]
                new_name_entry.delete(0, tk.END)
                new_name_entry.insert(0, data['name'])
                unit_name = self.get_unit_name_by_id(data['unit_id'])
                if unit_name:
                    unit_combo.set(unit_name)
                price_entry.delete(0, tk.END)
                price_entry.insert(0, str(data['price']))
            else:
                new_name_entry.delete(0, tk.END)
                price_entry.delete(0, tk.END)
                unit_combo.set('')

        existing_combo.bind('<<ComboboxSelected>>', on_existing_select)
        type_combo.bind('<<ComboboxSelected>>', lambda e: self.update_existing_items_list(type_combo.get(), existing_combo))

        self.update_existing_items_list('Материал', existing_combo)

        #  ТАБЛИЦА ПОЗИЦИЙ 
        columns = ('item_type', 'name', 'unit', 'quantity', 'accounting_quantity', 'price', 'total', 'location', 'responsible_person')
        items_tree = ttk.Treeview(items_frame, columns=columns, show='headings', height=8)

        headings = {'item_type': 'Тип', 'name': 'Наименование', 'unit': 'Ед.изм',
                    'quantity': 'Кол-во (факт)', 'accounting_quantity': 'Кол-во (учёт)',
                    'price': 'Цена', 'total': 'Сумма', 'location': 'Место', 'responsible_person': 'МОЛ'}
        for col, heading in headings.items():
            items_tree.heading(col, text=heading)
            items_tree.column(col, width=100 if col in ['item_type', 'quantity', 'accounting_quantity', 'price', 'total'] else 150)

        items_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        #  КНОПКИ УПРАВЛЕНИЯ ПОЗИЦИЯМИ 
        btn_items_frame = tk.Frame(items_frame, bg='white')
        btn_items_frame.pack(pady=5)

        def add_item():
            item_type = type_combo.get()
            selected_existing = existing_combo.get()
            new_name = new_name_entry.get().strip()
            unit = unit_combo.get().strip()
            location = location_combo.get().strip()
            responsible_person = responsible_person_combo.get().strip()

            try:
                quantity = float(quantity_entry.get()) if quantity_entry.get() else 0
                accounting_quantity = float(accounting_quantity_entry.get()) if accounting_quantity_entry.get() else quantity
                price = float(price_entry.get()) if price_entry.get() else 0
                inv_number = inv_number_entry.get().strip()
            except ValueError:
                messagebox.showerror("Ошибка", "Количество и цена должны быть числами")
                return

            if quantity <= 0:
                messagebox.showerror("Ошибка", "Количество должно быть больше 0")
                return

            if selected_existing != "➕ Новое" and selected_existing in self.existing_items_data:
                name = self.existing_items_data[selected_existing]['name']
            elif new_name:
                name = new_name
            else:
                messagebox.showerror("Ошибка", "Выберите существующую позицию или введите новое наименование")
                return

            if item_type == 'Основное средство' and inv_number:
                name = f"{name} (инв.№{inv_number})"

            total = quantity * price

            items_tree.insert('', tk.END, values=(
                item_type, name, unit, quantity, accounting_quantity, price, total, location, responsible_person
            ))

            if selected_existing == "➕ Новое":
                new_name_entry.delete(0, tk.END)
                price_entry.delete(0, tk.END)
                unit_combo.set('')
                inv_number_entry.delete(0, tk.END)
                location_combo.set('')
                responsible_person_combo.set('')
            quantity_entry.delete(0, tk.END)
            accounting_quantity_entry.delete(0, tk.END)
            quantity_entry.focus_set()

            update_total()

        tk.Button(btn_items_frame, text="➕ Добавить позицию", command=add_item,
                  bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_items_frame, text="🗑️ Удалить выбранную",
                  command=lambda: self.delete_selected_tree_item(items_tree),
                  bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        #  ИТОГОВАЯ СУММА 
        total_frame = tk.Frame(main_frame, bg='white')
        total_frame.pack(fill=tk.X, pady=10, padx=10)
        tk.Label(total_frame, text="ИТОГО:", font=('Arial', 14, 'bold'), bg='white', fg='#e74c3c').pack(side=tk.RIGHT, padx=10)
        total_label = tk.Label(total_frame, text="0.00 руб.", font=('Arial', 14, 'bold'), bg='white', fg='#e74c3c')
        total_label.pack(side=tk.RIGHT, padx=10)

        def update_total():
            total = 0
            for item in items_tree.get_children():
                values = items_tree.item(item)['values']
                if len(values) >= 7 and values[6]:
                    try:
                        total += float(values[6])
                    except:
                        pass
            total_label.config(text=f"{total:,.2f} руб.")
            return total

        #  ФУНКЦИЯ СОХРАНЕНИЯ 
        def save_act():
            try:
                act_date = date_entry.get().strip()
                supplier_name = supplier_combo.get().strip()
                receipt_type = 'purchase' if receipt_type_combo.get() == 'покупка' else 'free'
                person_name = person_combo.get().strip()

                if not supplier_name:
                    messagebox.showerror("Ошибка", "Введите поставщика")
                    return

                if items_tree.get_children() == ():
                    messagebox.showerror("Ошибка", "Добавьте хотя бы одну позицию")
                    return

                total_sum = update_total()

                conn2 = get_db_connection()
                cur2 = conn2.cursor()

                cur2.execute("SELECT id FROM suppliers WHERE name = %s", (supplier_name,))
                sup = cur2.fetchone()
                if sup:
                    supplier_id = sup[0]
                else:
                    cur2.execute("INSERT INTO suppliers (name) VALUES (%s) RETURNING id", (supplier_name,))
                    supplier_id = cur2.fetchone()[0]

                person_id = None
                if person_name:
                    cur2.execute("SELECT id FROM responsible_persons WHERE full_name = %s", (person_name,))
                    pers = cur2.fetchone()
                    if pers:
                        person_id = pers[0]
                    else:
                        cur2.execute("INSERT INTO responsible_persons (full_name) VALUES (%s) RETURNING id", (person_name,))
                        person_id = cur2.fetchone()[0]

                cur2.execute("""
                    INSERT INTO receipt_acts (act_number, act_date, supplier_id, supplier_name, receipt_type, 
                                              responsible_person_id, total_amount, status, created_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, 'draft', %s)
                    RETURNING id
                """, (act_number, act_date, supplier_id, supplier_name, receipt_type, person_id, total_sum, self.username))
                act_id = cur2.fetchone()[0]

                for item in items_tree.get_children():
                    values = items_tree.item(item)['values']
                    item_type = values[0]
                    name = values[1]
                    unit = values[2]
                    quantity = float(values[3]) if values[3] else 0
                    accounting_quantity = float(values[4]) if values[4] else quantity
                    price = float(values[5]) if values[5] else 0
                    total = quantity * price
                    location = values[7] if len(values) > 7 and values[7] else ''
                    responsible_person = values[8] if len(values) > 8 and values[8] else ''

                    unit_id = None
                    if unit:
                        cur2.execute("SELECT id FROM units WHERE short_name = %s OR full_name = %s", (unit, unit))
                        u = cur2.fetchone()
                        if u:
                            unit_id = u[0]

                    cur2.execute("""
                        INSERT INTO receipt_act_items (receipt_act_id, item_type, name, unit_id, quantity, 
                                                      accounting_quantity, price, total, location, responsible_person)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (act_id, 'material' if item_type == 'Материал' else 'inventory',
                          name, unit_id, quantity, accounting_quantity, price, total, location, responsible_person))

                conn2.commit()
                cur2.close()
                conn2.close()

                messagebox.showinfo("Успех", f"Акт {act_number} создан в статусе 'черновик'")
                # Проверяем наличие расхождений и показываем уведомление
                self.show_discrepancy_notification()
                dialog.destroy()

                if hasattr(self, 'receipt_acts_tree') and self.receipt_acts_tree:
                    self.load_receipt_acts()

                if messagebox.askyesno("Экспорт", "Создать акт в Word/Excel?"):
                    self.export_receipt_act_to_word_by_id(act_id)
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        #  КНОПКИ СОХРАНЕНИЯ (В КОНЦЕ main_frame) 
        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=20, fill=tk.X, padx=10)
        tk.Button(btn_frame, text="💾 Сохранить как черновик", command=save_act,
                  bg='#f39c12', fg='white', width=25, height=2).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="✖️ Отмена", command=dialog.destroy,
                  bg='#95a5a6', fg='white', width=20, height=2).pack(side=tk.LEFT, padx=10)

    def create_write_off_act(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Создание акта списания")
        dialog.geometry("900x750")
        dialog.configure(bg='white')
        dialog.transient(self.root)  # Связываем с главным окном
        dialog.grab_set()  # Делаем модальным (блокируем другие окна)
        dialog.focus_set()  # Устанавливаем фокус

        #  ОСНОВНОЙ КОНТЕЙНЕР 
        main_container = tk.Frame(dialog, bg='white')
        main_container.pack(fill=tk.BOTH, expand=True)

        #  СКРОЛЛИНГ 
        canvas = tk.Canvas(main_container, bg='white')
        scrollbar = tk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = scrollable_frame

        tk.Label(main_frame, text="АКТ СПИСАНИЯ", font=('Arial', 16, 'bold'), bg='white').pack(pady=10)

        #  ИНФОРМАЦИЯ ОБ АКТЕ 
        info_frame = tk.LabelFrame(main_frame, text="Информация об акте", font=('Arial', 12, 'bold'), bg='white')
        info_frame.pack(fill=tk.X, pady=10, padx=10)

        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COALESCE(MAX(id), 0) + 1 FROM write_off_acts")
        next_id = cur.fetchone()[0]
        cur.close()
        conn.close()
        act_number = f"АС-{next_id:04d}"

        tk.Label(info_frame, text="Номер акта:", font=('Arial', 10), bg='white').grid(row=0, column=0, padx=10, pady=5,
                                                                                      sticky='w')
        tk.Label(info_frame, text=act_number, font=('Arial', 10, 'bold'), bg='white', fg='#27ae60').grid(row=0,
                                                                                                         column=1,
                                                                                                         padx=10,
                                                                                                         pady=5,
                                                                                                         sticky='w')

        tk.Label(info_frame, text="Дата акта:", font=('Arial', 10), bg='white').grid(row=0, column=2, padx=10, pady=5,
                                                                                     sticky='w')
        date_entry = tk.Entry(info_frame, width=15)
        date_entry.insert(0, datetime.now().strftime('%Y-%m-%d'))
        date_entry.grid(row=0, column=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Тип списания:", font=('Arial', 10), bg='white').grid(row=1, column=0, padx=10,
                                                                                        pady=5, sticky='w')
        type_combo = ttk.Combobox(info_frame, values=['материалы', 'основные_средства'], width=40)
        type_combo.current(0)
        type_combo.grid(row=1, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        tk.Label(info_frame, text="Причина списания:", font=('Arial', 10), bg='white').grid(row=2, column=0, padx=10,
                                                                                            pady=5, sticky='w')
        reason_entry = tk.Text(info_frame, height=3, width=40)
        reason_entry.grid(row=2, column=1, columnspan=3, padx=10, pady=5, sticky='w')

        # Комиссия
        commission = self.get_commission()
        commission_frame = tk.LabelFrame(info_frame, text="Состав комиссии", font=('Arial', 10, 'bold'), bg='white',
                                         fg='#2c3e50')
        commission_frame.grid(row=3, column=0, columnspan=4, padx=10, pady=5, sticky='we')

        if commission:
            # Отображаем председателя
            tk.Label(commission_frame, text="Председатель:", font=('Arial', 9, 'bold'), bg='white', fg='#27ae60').grid(
                row=0, column=0, padx=10, pady=2, sticky='w')
            tk.Label(commission_frame, text=commission[1], font=('Arial', 9), bg='white', wraplength=500).grid(row=0,
                                                                                                               column=1,
                                                                                                               padx=10,
                                                                                                               pady=2,
                                                                                                               sticky='w')

            # Отображаем членов комиссии
            tk.Label(commission_frame, text="Члены комиссии:", font=('Arial', 9, 'bold'), bg='white',
                     fg='#27ae60').grid(row=1, column=0, padx=10, pady=2, sticky='w')
            tk.Label(commission_frame, text=commission[2], font=('Arial', 9), bg='white', wraplength=500).grid(row=1,
                                                                                                               column=1,
                                                                                                               padx=10,
                                                                                                               pady=2,
                                                                                                               sticky='w')
        else:
            tk.Label(commission_frame, text="Комиссия не назначена. Обратитесь к администратору.",
                     font=('Arial', 9, 'italic'), bg='white', fg='red').pack(pady=5, padx=10)

        #  ПОЗИЦИИ СПИСАНИЯ 
        items_frame = tk.LabelFrame(main_frame, text="Позиции списания", font=('Arial', 12, 'bold'), bg='white')
        items_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)

        # Форма для добавления позиции (упрощённая)
        add_frame = tk.Frame(items_frame, bg='white', relief=tk.GROOVE, bd=2)
        add_frame.pack(fill=tk.X, padx=10, pady=10)

        tk.Label(add_frame, text="Тип:", font=('Arial', 10, 'bold'), bg='white').grid(row=0, column=0, padx=5, pady=5,
                                                                                      sticky='w')
        write_off_type_combo = ttk.Combobox(add_frame, values=['материалы', 'основные_средства'], width=15)
        write_off_type_combo.current(0)
        write_off_type_combo.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(add_frame, text="Выбрать позицию:", font=('Arial', 10, 'bold'), bg='white').grid(row=0, column=2,
                                                                                                  padx=5, pady=5,
                                                                                                  sticky='w')
        existing_combo = ttk.Combobox(add_frame, width=40)
        existing_combo.grid(row=0, column=3, padx=5, pady=5)

        tk.Button(add_frame, text="Обновить",
                  command=lambda: self.update_write_off_items_list(write_off_type_combo.get(), existing_combo),
                  bg='#3498db', fg='white', width=8).grid(row=0, column=4, padx=2, pady=5)

        tk.Label(add_frame, text="Количество:", font=('Arial', 10, 'bold'), bg='white').grid(row=1, column=0, padx=5,
                                                                                             pady=5, sticky='w')
        quantity_entry = tk.Entry(add_frame, width=15)
        quantity_entry.grid(row=1, column=1, padx=5, pady=5)

        tk.Label(add_frame, text="Цена (автоматически):", font=('Arial', 10, 'bold'), bg='white').grid(row=1, column=2,
                                                                                                       padx=5, pady=5,
                                                                                                       sticky='w')
        price_label = tk.Label(add_frame, text="0.00 руб.", font=('Arial', 10), bg='white', fg='#e74c3c', width=15,
                               anchor='w')
        price_label.grid(row=1, column=3, padx=5, pady=5, sticky='w')

        tk.Label(add_frame, text="Ед.изм (автоматически):", font=('Arial', 10, 'bold'), bg='white').grid(row=2,
                                                                                                         column=0,
                                                                                                         padx=5, pady=5,
                                                                                                         sticky='w')
        unit_label = tk.Label(add_frame, text="", font=('Arial', 10), bg='white', width=15, anchor='w')
        unit_label.grid(row=2, column=1, padx=5, pady=5, sticky='w')

        #  ФУНКЦИЯ ОБНОВЛЕНИЯ СПИСКА 
        def on_existing_select(event):
            selected = existing_combo.get()
            if selected and selected in self.write_off_items_data:
                data = self.write_off_items_data[selected]
                unit_name = self.get_unit_name_by_id(data['unit_id'])
                unit_label.config(text=unit_name if unit_name else '-')
                price_label.config(text=f"{data['price']:,.2f} руб.")
            else:
                unit_label.config(text='')
                price_label.config(text='0.00 руб.')

        existing_combo.bind('<<ComboboxSelected>>', on_existing_select)
        write_off_type_combo.bind('<<ComboboxSelected>>',
                                  lambda e: self.update_write_off_items_list(write_off_type_combo.get(),
                                                                             existing_combo))

        self.update_write_off_items_list('материалы', existing_combo)

        #  ТАБЛИЦА ПОЗИЦИЙ 
        columns = ('name', 'unit', 'quantity', 'price', 'total')
        items_tree = ttk.Treeview(items_frame, columns=columns, show='headings', height=10)

        headings = {'name': 'Наименование', 'unit': 'Ед.изм', 'quantity': 'Кол-во', 'price': 'Цена', 'total': 'Сумма'}
        for col, heading in headings.items():
            items_tree.heading(col, text=heading)
            items_tree.column(col, width=150 if col == 'name' else 100)

        items_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        #  КНОПКИ УПРАВЛЕНИЯ ПОЗИЦИЯМИ 
        btn_items_frame = tk.Frame(items_frame, bg='white')
        btn_items_frame.pack(pady=5)

        def add_item():
            selected_existing = existing_combo.get()

            if not selected_existing or selected_existing not in self.write_off_items_data:
                messagebox.showerror("Ошибка", "Выберите позицию для списания")
                return

            try:
                quantity = float(quantity_entry.get()) if quantity_entry.get() else 0
            except ValueError:
                messagebox.showerror("Ошибка", "Количество должно быть числом")
                return

            if quantity <= 0:
                messagebox.showerror("Ошибка", "Количество должно быть больше 0")
                return

            data = self.write_off_items_data[selected_existing]
            if quantity > data['actual_quantity']:
                messagebox.showerror("Ошибка", f"Недостаточно на складе! Доступно: {data['actual_quantity']} шт.")
                return

            name = data['name']
            unit = self.get_unit_name_by_id(data['unit_id'])
            # Преобразуем Decimal в float для умножения
            price = float(data['price']) if data['price'] else 0
            total = quantity * price

            items_tree.insert('', tk.END, values=(name, unit, quantity, price, total))

            quantity_entry.delete(0, tk.END)
            quantity_entry.focus_set()

            update_total()

        tk.Button(btn_items_frame, text="➕ Добавить позицию", command=add_item,
                  bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_items_frame, text="🗑️ Удалить выбранную",
                  command=lambda: self.delete_selected_tree_item(items_tree),
                  bg='#e74c3c', fg='white').pack(side=tk.LEFT, padx=5)

        #  ИТОГОВАЯ СУММА 
        total_frame = tk.Frame(main_frame, bg='white')
        total_frame.pack(fill=tk.X, pady=10, padx=10)
        tk.Label(total_frame, text="ИТОГО:", font=('Arial', 14, 'bold'), bg='white', fg='#e74c3c').pack(side=tk.RIGHT,
                                                                                                        padx=10)
        total_label = tk.Label(total_frame, text="0.00 руб.", font=('Arial', 14, 'bold'), bg='white', fg='#e74c3c')
        total_label.pack(side=tk.RIGHT, padx=10)

        def update_total():
            total = 0
            for item in items_tree.get_children():
                values = items_tree.item(item)['values']
                if len(values) >= 5 and values[4]:
                    try:
                        total += float(values[4])
                    except:
                        pass
            total_label.config(text=f"{total:,.2f} руб.")
            return total

        #  ФУНКЦИЯ СОХРАНЕНИЯ 
        def save_act():
            try:
                act_date = date_entry.get().strip()
                write_off_type = type_combo.get()
                reason = reason_entry.get('1.0', tk.END).strip()

                if items_tree.get_children() == ():
                    messagebox.showerror("Ошибка", "Добавьте хотя бы одну позицию")
                    return

                if not commission:
                    if not messagebox.askyesno("Предупреждение", "Комиссия не назначена. Продолжить?"):
                        return

                total_sum = update_total()

                conn2 = get_db_connection()
                cur2 = conn2.cursor()

                commission_id = commission[0] if commission else None

                cur2.execute("""
                    INSERT INTO write_off_acts (act_number, act_date, write_off_type, total_amount, status, reason, created_by, commission_id)
                    VALUES (%s, %s, %s, %s, 'draft', %s, %s, %s)
                    RETURNING id
                """, (act_number, act_date, write_off_type, total_sum, reason, self.username, commission_id))
                act_id = cur2.fetchone()[0]

                for item in items_tree.get_children():
                    values = items_tree.item(item)['values']
                    name = values[0]
                    unit = values[1]
                    quantity = float(values[2]) if values[2] else 0
                    price = float(values[3]) if values[3] else 0
                    total = quantity * price

                    unit_id = None
                    if unit:
                        cur2.execute("SELECT id FROM units WHERE short_name = %s OR full_name = %s", (unit, unit))
                        u = cur2.fetchone()
                        if u:
                            unit_id = u[0]

                    cur2.execute("""
                        INSERT INTO write_off_act_items (write_off_act_id, item_type, name, unit_id, quantity, price, total, reason)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """, (act_id, write_off_type, name, unit_id, quantity, price, total, reason))

                conn2.commit()
                cur2.close()
                conn2.close()

                messagebox.showinfo("Успех", f"Акт {act_number} создан в статусе 'черновик'")

                self.show_discrepancy_notification()

                dialog.destroy()
                self.load_write_off_acts()
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

        #  КНОПКИ СОХРАНЕНИЯ 
        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=20, fill=tk.X, padx=10)
        tk.Button(btn_frame, text="💾 Сохранить как черновик", command=save_act,
                  bg='#f39c12', fg='white', width=25, height=2).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="✖️ Отмена", command=dialog.destroy,
                  bg='#95a5a6', fg='white', width=20, height=2).pack(side=tk.LEFT, padx=10)

    def logout(self):
        if messagebox.askyesno("Выход", "Вы уверены?"):
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO action_logs (username, action, object_type, details, created_at)
                    VALUES (%s, %s, %s, %s, %s)
                """, (self.username, 'LOGOUT', 'user',
                      f'Пользователь {self.username} ({self.full_name}) вышел из системы', datetime.now()))
                conn.commit()
                cur.close()
                conn.close()
            except:
                pass
            self.root.destroy()
            LoginApp()


class LoginApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Авторизация - Школьный склад")
        self.root.geometry("400x300")
        self.root.resizable(False, False)

        tk.Label(self.root, text="Вход в систему", font=('Arial', 18, 'bold')).pack(pady=30)

        tk.Label(self.root, text="Логин:", font=('Arial', 12)).pack()
        self.username_entry = tk.Entry(self.root, font=('Arial', 12), width=30)
        self.username_entry.pack(pady=5)

        tk.Label(self.root, text="Пароль:", font=('Arial', 12)).pack()
        self.password_entry = tk.Entry(self.root, font=('Arial', 12), width=30, show='*')
        self.password_entry.pack(pady=5)

        self.error_label = tk.Label(self.root, text="", fg='red')
        self.error_label.pack(pady=5)

        tk.Button(self.root, text="Войти", command=self.login, font=('Arial', 12), bg='#27ae60', fg='white',
                  width=20).pack(pady=20)

        self.root.bind('<Return>', lambda e: self.login())

        self.root.mainloop()

    def login(self):
        username = self.username_entry.get().strip()
        password = self.password_entry.get().strip()
        password_hash = hash_password(password)

        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("""
                SELECT u.id, u.role, u.person_id, rp.full_name 
                FROM users u
                LEFT JOIN responsible_persons rp ON u.person_id = rp.id
                WHERE u.username = %s AND u.password_hash = %s AND u.is_active = TRUE
            """, (username, password_hash))
            row = cur.fetchone()

            if row:
                user_id, role, person_id, full_name = row

                cur.execute("""
                    INSERT INTO action_logs (username, action, object_type, details, created_at)
                    VALUES (%s, %s, %s, %s, %s)
                """, (username, 'LOGIN', 'user', f'Пользователь {username} ({full_name}) вошёл в систему',
                      datetime.now()))
                conn.commit()

                cur.close()
                conn.close()

                self.root.destroy()
                MainApp(user_id, username, role, full_name)
            else:
                self.error_label.config(text="Неверный логин или пароль")
                cur.close()
                conn.close()
        except Exception as e:
            self.error_label.config(text=f"Ошибка подключения: {e}")


if __name__ == "__main__":
    LoginApp()